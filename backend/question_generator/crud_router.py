from fastapi import (
    APIRouter,
    Depends,
    HTTPException
)

from fastapi.responses import FileResponse

from sqlalchemy.orm import Session

from backend.database import get_db

from backend.models import (
    QuestionSet,
    Question
)

from .schemas import (
    QuestionResponse,
    QuestionSetResponse,
    QuestionUpdateRequest
)

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import getSampleStyleSheet

from docx import Document



router = APIRouter(
    prefix="/question-management",
    tags=["Question Management"]
)



# ======================================
# SAVE QUESTION SET
# Instructor Approved Save
# ======================================

@router.post("/save")
def save_question_set(

    data: dict,

    db: Session = Depends(get_db)

):


    question_set = QuestionSet(

        title=data.get(
            "title",
            "Generated Assessment"
        ),

        source_type=data.get(
            "source_type",
            "topic"
        ),

        topic=data.get(
            "topic"
        ),

        difficulty=data.get(
            "difficulty"
        ),

        total_questions=data.get(
            "total_questions",
            0
        ),

        total_marks=data.get(
            "total_marks",
            0
        )

    )


    db.add(
        question_set
    )

    db.commit()

    db.refresh(
        question_set
    )



    for q in data.get(
        "questions",
        []
    ):


        question = Question(

            question_set_id=question_set.id,

            question_number=q.get(
                "question_number"
            ),

            type=q.get(
                "type"
            ),

            question_text=q.get(
                "question"
            ),

            options=q.get(
                "options"
            ),

            correct_answer=q.get(
                "correct_answer"
            ),

            marks=q.get(
                "marks"
            )

        )


        db.add(
            question
        )



    db.commit()



    return {

        "message":
        "Question set saved successfully",

        "question_set_id":
        question_set.id

    }





# ======================================
# GET ALL QUESTION SETS
# ======================================

@router.get(
    "/sets",
    response_model=list[QuestionSetResponse]
)
def get_question_sets(

    db: Session = Depends(get_db)

):

    return db.query(
        QuestionSet
    ).all()




# ======================================
# GET QUESTION SET DETAILS
# ======================================

@router.get(
    "/sets/{set_id}"
)
def get_question_set(

    set_id: int,

    db: Session = Depends(get_db)

):

    question_set = db.query(
        QuestionSet
    ).filter(
        QuestionSet.id == set_id
    ).first()



    if not question_set:

        raise HTTPException(

            status_code=404,

            detail="Question set not found"

        )



    questions = db.query(
        Question
    ).filter(
        Question.question_set_id == set_id
    ).all()



    return {

        "question_set": question_set,

        "questions": questions

    }





# ======================================
# UPDATE QUESTION
# ======================================

@router.put(
    "/questions/{question_id}",
    response_model=QuestionResponse
)
def update_question(

    question_id: int,

    question_data: QuestionUpdateRequest,

    db: Session = Depends(get_db)

):


    question = db.query(
        Question
    ).filter(
        Question.id == question_id
    ).first()



    if not question:

        raise HTTPException(

            status_code=404,

            detail="Question not found"

        )



    for key, value in question_data.model_dump(
        exclude_unset=True
    ).items():

        setattr(
            question,
            key,
            value
        )



    db.commit()

    db.refresh(
        question
    )



    return question





# ======================================
# EXPORT PDF
# ======================================

@router.get(
    "/sets/{set_id}/export/pdf"
)
def export_pdf(

    set_id: int,

    db: Session = Depends(get_db)

):

    question_set = db.query(
        QuestionSet
    ).filter(
        QuestionSet.id == set_id
    ).first()



    if not question_set:

        raise HTTPException(

            status_code=404,

            detail="Question set not found"

        )



    questions = db.query(
        Question
    ).filter(
        Question.question_set_id == set_id
    ).all()



    file_name = (
        f"question_set_{set_id}.pdf"
    )


    doc = SimpleDocTemplate(
        file_name
    )


    styles = getSampleStyleSheet()

    content = []



    content.append(
        Paragraph(
            question_set.title,
            styles["Title"]
        )
    )



    content.append(
        Spacer(1,20)
    )



    for q in questions:


        content.append(
            Paragraph(

                f"Q{q.question_number}. "
                f"{q.question_text}<br/>"
                f"Type: {q.type}<br/>"
                f"Marks: {q.marks}<br/>"
                f"Answer: {q.correct_answer}",

                styles["Normal"]

            )
        )


        content.append(
            Spacer(1,15)
        )



    doc.build(
        content
    )



    return FileResponse(

        file_name,

        filename=file_name,

        media_type="application/pdf"

    )





# ======================================
# EXPORT DOCX
# ======================================

@router.get(
    "/sets/{set_id}/export/docx"
)
def export_docx(

    set_id: int,

    db: Session = Depends(get_db)

):


    question_set = db.query(
        QuestionSet
    ).filter(
        QuestionSet.id == set_id
    ).first()



    if not question_set:

        raise HTTPException(

            status_code=404,

            detail="Question set not found"

        )



    questions = db.query(
        Question
    ).filter(
        Question.question_set_id == set_id
    ).all()



    file_name = (
        f"question_set_{set_id}.docx"
    )


    doc = Document()



    doc.add_heading(
        question_set.title,
        level=1
    )



    for q in questions:

        doc.add_paragraph(

            f"Q{q.question_number}. "
            f"{q.question_text}\n"
            f"Type: {q.type}\n"
            f"Marks: {q.marks}\n"
            f"Answer: {q.correct_answer}"

        )



    doc.save(
        file_name
    )



    return FileResponse(

        file_name,

        filename=file_name,

        media_type=(
            "application/vnd.openxmlformats-"
            "officedocument.wordprocessingml.document"
        )

    )





# ======================================
# DELETE QUESTION SET
# ======================================

@router.delete(
    "/sets/{set_id}"
)
def delete_question_set(

    set_id: int,

    db: Session = Depends(get_db)

):

    question_set = db.query(
        QuestionSet
    ).filter(
        QuestionSet.id == set_id
    ).first()



    if not question_set:

        raise HTTPException(

            status_code=404,

            detail="Question set not found"

        )



    db.delete(
        question_set
    )

    db.commit()



    return {

        "message":
        "Question set deleted successfully"

    }