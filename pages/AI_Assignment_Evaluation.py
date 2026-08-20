import hashlib
import base64
import html
import io
import os
import re
import sys
import zipfile
from datetime import datetime
import json

import streamlit as st


# =========================================================
# PROJECT ROOT
# =========================================================

PROJECT_ROOT = os.path.abspath(
    os.path.join(
        os.path.dirname(__file__),
        "..",
    )
)

if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)


from api_client import (  # noqa: E402
    check_assignment_evaluation_health,
    evaluate_assignment,
    generate_assignment_rubric,
    get_question_sets,
    get_question_set,
    download_pdf,
    save_batch_evaluation_results,
)


# =========================================================
# PAGE CONFIG
# =========================================================

st.set_page_config(
    page_title="AI Assignment Evaluation",
    layout="wide",
    initial_sidebar_state="collapsed",
)


# =========================================================
# SESSION STATE
# =========================================================

DEFAULT_STATE = {
    "ai_evaluation_result": None,
    "ai_evaluation_batch_results": [],
    "ai_evaluation_error": None,
    "ai_evaluation_snapshot": None,
    "ai_evaluation_phase": "idle",
    "ai_evaluation_logs": [],
    "rubric_assignment_hash": None,
    "rubric_assignment_snapshot": None,
    "generated_rubric": None,
    "rubric_generation_phase": "idle",
    "rubric_generation_error": None,
    "rubric_source": None,
    "uploaded_rubric_snapshot": None,
    "show_rubric_upload": False,
    "open_rubric_edit_after_generation": False,
}


for key, value in DEFAULT_STATE.items():
    if key not in st.session_state:
        st.session_state[key] = value


# =========================================================
# HELPERS
# =========================================================

def safe_number(value, default=0):
    try:
        return float(value)
    except (TypeError, ValueError):
        return float(default)


def format_score(value):
    number = safe_number(value)
    if number.is_integer():
        return str(int(number))
    return f"{number:.2f}"


def safe_html(value):
    if value is None:
        return ""
    return html.escape(str(value))


def get_evaluation_payload(api_result):
    if not isinstance(api_result, dict):
        return {}
    evaluation = api_result.get("evaluation")
    return evaluation if isinstance(evaluation, dict) else {}


def build_evaluation_report_text(
    evaluation,
    snapshot,
):
    """
    Build an instructor-friendly downloadable marks report entirely
    in memory. No extra project file or backend endpoint is required.
    """

    assignment_name = (
        evaluation.get("assignment_name")
        or "Assignment Evaluation"
    )

    total_marks = safe_number(
        evaluation.get(
            "max_score",
            evaluation.get(
                "total_marks",
                0,
            ),
        )
    )

    obtained_marks = safe_number(
        evaluation.get(
            "total_score",
            evaluation.get(
                "obtained_marks",
                0,
            ),
        )
    )

    percentage = (
        (obtained_marks / total_marks) * 100
        if total_marks > 0
        else 0.0
    )

    remarks = (
        evaluation.get("overall_feedback")
        or evaluation.get("remarks")
        or "No overall remarks provided."
    )

    assignment_file = (
        (snapshot or {})
        .get("assignment_file", {})
        .get("name", "-")
    )

    rubric_file = (
        (snapshot or {})
        .get("rubric_file", {})
        .get("name", "-")
    )

    submission_file = (
        (snapshot or {})
        .get("submission_file", {})
        .get("name", "-")
    )

    lines = [
        "EduAI - AI Assignment Evaluation Report",
        "=" * 48,
        "",
        f"Assignment: {assignment_name}",
        f"Assignment File: {assignment_file}",
        f"Rubric File: {rubric_file}",
        f"Student Submission: {submission_file}",
        f"Evaluated At: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "FINAL MARKS",
        "-" * 48,
        f"Total Marks: {format_score(total_marks)}",
        f"Marks Obtained: {format_score(obtained_marks)}",
        f"Percentage: {percentage:.2f}%",
        "",
        "REMARKS",
        "-" * 48,
        str(remarks),
        "",
    ]

    criteria = evaluation.get("criteria", [])
    deductions = evaluation.get("deductions", [])

    if criteria:
        lines.extend(
            [
                "CRITERION-WISE EVALUATION",
                "-" * 48,
            ]
        )

        for index, criterion in enumerate(
            criteria,
            start=1,
        ):
            name = (
                criterion.get("criterion")
                or criterion.get("name")
                or f"Criterion {index}"
            )

            score = safe_number(
                criterion.get(
                    "score",
                    0,
                )
            )

            max_score = safe_number(
                criterion.get(
                    "max_score",
                    criterion.get(
                        "max_marks",
                        0,
                    ),
                )
            )

            evidence = (
                criterion.get("evidence")
                or "-"
            )

            feedback = (
                criterion.get("feedback")
                or "-"
            )

            lines.extend(
                [
                    f"{index}. {name}",
                    (
                        "   Score: "
                        f"{format_score(score)} / "
                        f"{format_score(max_score)}"
                    ),
                    f"   Evidence: {evidence}",
                    f"   Feedback: {feedback}",
                    "",
                ]
            )

    elif deductions:
        lines.extend(
            [
                "MARK DEDUCTIONS",
                "-" * 48,
            ]
        )

        for index, deduction in enumerate(
            deductions,
            start=1,
        ):
            question = (
                deduction.get("question")
                or f"Item {index}"
            )

            marks_deducted = safe_number(
                deduction.get(
                    "marks_deducted",
                    0,
                )
            )

            reason = (
                deduction.get("reason")
                or "-"
            )

            lines.extend(
                [
                    f"{index}. {question}",
                    (
                        "   Marks Deducted: "
                        f"{format_score(marks_deducted)}"
                    ),
                    f"   Reason: {reason}",
                    "",
                ]
            )

    lines.extend(
        [
            "-" * 48,
            "Generated by EduAI Assignment Evaluation",
        ]
    )

    return "\n".join(lines)


def evaluation_report_filename(
    snapshot,
    extension="txt",
):
    submission_name = (
        (snapshot or {})
        .get("submission_file", {})
        .get("name", "student_submission")
    )

    base_name = os.path.splitext(
        submission_name
    )[0]

    safe_name = re.sub(
        r"[^A-Za-z0-9_-]+",
        "_",
        base_name,
    ).strip("_")

    if not safe_name:
        safe_name = "student_submission"

    extension = str(
        extension
    ).lower().lstrip(".")

    return f"{safe_name}_evaluation_report.{extension}"


def _report_common_values(
    evaluation,
    snapshot,
):
    assignment_name = (
        evaluation.get("assignment_name")
        or "Assignment Evaluation"
    )

    total_marks = safe_number(
        evaluation.get(
            "max_score",
            evaluation.get(
                "total_marks",
                0,
            ),
        )
    )

    obtained_marks = safe_number(
        evaluation.get(
            "total_score",
            evaluation.get(
                "obtained_marks",
                0,
            ),
        )
    )

    percentage = (
        (obtained_marks / total_marks) * 100
        if total_marks > 0
        else 0.0
    )

    remarks = (
        evaluation.get("overall_feedback")
        or evaluation.get("remarks")
        or "No overall remarks provided."
    )

    assignment_file = (
        (snapshot or {})
        .get("assignment_file", {})
        .get("name", "-")
    )

    rubric_file = (
        (snapshot or {})
        .get("rubric_file", {})
        .get("name", "-")
    )

    submission_file = (
        (snapshot or {})
        .get("submission_file", {})
        .get("name", "-")
    )

    return {
        "assignment_name": assignment_name,
        "total_marks": total_marks,
        "obtained_marks": obtained_marks,
        "percentage": percentage,
        "remarks": str(remarks),
        "assignment_file": assignment_file,
        "rubric_file": rubric_file,
        "submission_file": submission_file,
    }


def build_evaluation_report_docx(
    evaluation,
    snapshot,
):
    """Build a Word (.docx) evaluation report in memory."""

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as error:
        raise RuntimeError(
            "Word export requires python-docx. "
            "Install it with: python -m pip install python-docx"
        ) from error

    values = _report_common_values(
        evaluation,
        snapshot,
    )

    buffer = io.BytesIO()
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(
        "EduAI - AI Assignment Evaluation Report"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)

    document.add_paragraph(
        f"Assignment: {values['assignment_name']}"
    )
    document.add_paragraph(
        f"Assignment File: {values['assignment_file']}"
    )
    document.add_paragraph(
        f"Rubric File: {values['rubric_file']}"
    )
    document.add_paragraph(
        f"Student Submission: {values['submission_file']}"
    )
    document.add_paragraph(
        f"Evaluated At: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )

    document.add_heading(
        "Final Marks",
        level=1,
    )

    marks_table = document.add_table(
        rows=4,
        cols=2,
    )
    marks_table.style = "Table Grid"

    marks_rows = [
        ("Total Marks", format_score(values["total_marks"])),
        ("Marks Obtained", format_score(values["obtained_marks"])),
        ("Percentage", f"{values['percentage']:.2f}%"),
        ("Remarks", values["remarks"]),
    ]

    for row, (label, value) in zip(
        marks_table.rows,
        marks_rows,
    ):
        row.cells[0].text = str(label)
        row.cells[1].text = str(value)

    criteria = evaluation.get(
        "criteria",
        [],
    )
    deductions = evaluation.get(
        "deductions",
        [],
    )

    if criteria:
        document.add_heading(
            "Criterion-wise Evaluation",
            level=1,
        )

        for index, criterion in enumerate(
            criteria,
            start=1,
        ):
            name = (
                criterion.get("criterion")
                or criterion.get("name")
                or f"Criterion {index}"
            )
            score = safe_number(
                criterion.get(
                    "score",
                    0,
                )
            )
            max_score = safe_number(
                criterion.get(
                    "max_score",
                    criterion.get(
                        "max_marks",
                        0,
                    ),
                )
            )
            evidence = (
                criterion.get("evidence")
                or "-"
            )
            feedback = (
                criterion.get("feedback")
                or "-"
            )

            heading = document.add_paragraph()
            heading_run = heading.add_run(
                f"{index}. {name}"
            )
            heading_run.bold = True

            document.add_paragraph(
                f"Score: {format_score(score)} / "
                f"{format_score(max_score)}"
            )
            document.add_paragraph(
                f"Evidence: {evidence}"
            )
            document.add_paragraph(
                f"Feedback: {feedback}"
            )

    elif deductions:
        document.add_heading(
            "Mark Deductions",
            level=1,
        )

        for index, deduction in enumerate(
            deductions,
            start=1,
        ):
            question = (
                deduction.get("question")
                or f"Item {index}"
            )
            marks_deducted = safe_number(
                deduction.get(
                    "marks_deducted",
                    0,
                )
            )
            reason = (
                deduction.get("reason")
                or "-"
            )

            heading = document.add_paragraph()
            heading_run = heading.add_run(
                f"{index}. {question}"
            )
            heading_run.bold = True

            document.add_paragraph(
                "Marks Deducted: "
                f"{format_score(marks_deducted)}"
            )
            document.add_paragraph(
                f"Reason: {reason}"
            )

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        "Generated by EduAI Assignment Evaluation"
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)

    document.save(
        buffer
    )

    return buffer.getvalue()


def build_evaluation_report_pdf(
    evaluation,
    snapshot,
):
    """Build a PDF evaluation report in memory using ReportLab."""

    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import (
            ParagraphStyle,
            getSampleStyleSheet,
        )
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError as error:
        raise RuntimeError(
            "PDF export requires reportlab. "
            "Install it with: python -m pip install reportlab"
        ) from error

    values = _report_common_values(
        evaluation,
        snapshot,
    )

    buffer = io.BytesIO()

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title="EduAI Assignment Evaluation Report",
        author="EduAI",
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "EduAITitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontSize=16,
        leading=20,
        spaceAfter=12,
    )

    section_style = ParagraphStyle(
        "EduAISection",
        parent=styles["Heading2"],
        fontSize=12,
        leading=15,
        spaceBefore=10,
        spaceAfter=6,
    )

    body_style = ParagraphStyle(
        "EduAIBody",
        parent=styles["BodyText"],
        fontSize=9.5,
        leading=13,
        spaceAfter=4,
    )

    small_style = ParagraphStyle(
        "EduAISmall",
        parent=styles["BodyText"],
        fontSize=8.5,
        leading=11,
    )

    story = [
        Paragraph(
            "EduAI - AI Assignment Evaluation Report",
            title_style,
        ),
        Paragraph(
            f"<b>Assignment:</b> {html.escape(str(values['assignment_name']))}",
            body_style,
        ),
        Paragraph(
            f"<b>Assignment File:</b> {html.escape(str(values['assignment_file']))}",
            body_style,
        ),
        Paragraph(
            f"<b>Rubric File:</b> {html.escape(str(values['rubric_file']))}",
            body_style,
        ),
        Paragraph(
            f"<b>Student Submission:</b> {html.escape(str(values['submission_file']))}",
            body_style,
        ),
        Paragraph(
            "<b>Evaluated At:</b> "
            + datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            body_style,
        ),
        Spacer(1, 6),
        Paragraph(
            "Final Marks",
            section_style,
        ),
    ]

    marks_data = [
        ["Total Marks", format_score(values["total_marks"])],
        ["Marks Obtained", format_score(values["obtained_marks"])],
        ["Percentage", f"{values['percentage']:.2f}%"],
        ["Remarks", Paragraph(
            html.escape(values["remarks"]),
            small_style,
        )],
    ]

    marks_table = Table(
        marks_data,
        colWidths=[
            42 * mm,
            123 * mm,
        ],
    )

    marks_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#CBD5E1")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F8FAFC")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )

    story.extend(
        [
            marks_table,
            Spacer(1, 8),
        ]
    )

    criteria = evaluation.get(
        "criteria",
        [],
    )
    deductions = evaluation.get(
        "deductions",
        [],
    )

    if criteria:
        story.append(
            Paragraph(
                "Criterion-wise Evaluation",
                section_style,
            )
        )

        for index, criterion in enumerate(
            criteria,
            start=1,
        ):
            name = (
                criterion.get("criterion")
                or criterion.get("name")
                or f"Criterion {index}"
            )
            score = safe_number(
                criterion.get(
                    "score",
                    0,
                )
            )
            max_score = safe_number(
                criterion.get(
                    "max_score",
                    criterion.get(
                        "max_marks",
                        0,
                    ),
                )
            )
            evidence = (
                criterion.get("evidence")
                or "-"
            )
            feedback = (
                criterion.get("feedback")
                or "-"
            )

            story.extend(
                [
                    Paragraph(
                        f"<b>{index}. {html.escape(str(name))}</b>",
                        body_style,
                    ),
                    Paragraph(
                        "Score: "
                        f"{format_score(score)} / "
                        f"{format_score(max_score)}",
                        body_style,
                    ),
                    Paragraph(
                        "<b>Evidence:</b> "
                        f"{html.escape(str(evidence))}",
                        body_style,
                    ),
                    Paragraph(
                        "<b>Feedback:</b> "
                        f"{html.escape(str(feedback))}",
                        body_style,
                    ),
                    Spacer(1, 5),
                ]
            )

    elif deductions:
        story.append(
            Paragraph(
                "Mark Deductions",
                section_style,
            )
        )

        for index, deduction in enumerate(
            deductions,
            start=1,
        ):
            question = (
                deduction.get("question")
                or f"Item {index}"
            )
            marks_deducted = safe_number(
                deduction.get(
                    "marks_deducted",
                    0,
                )
            )
            reason = (
                deduction.get("reason")
                or "-"
            )

            story.extend(
                [
                    Paragraph(
                        f"<b>{index}. {html.escape(str(question))}</b>",
                        body_style,
                    ),
                    Paragraph(
                        "Marks Deducted: "
                        f"{format_score(marks_deducted)}",
                        body_style,
                    ),
                    Paragraph(
                        "<b>Reason:</b> "
                        f"{html.escape(str(reason))}",
                        body_style,
                    ),
                    Spacer(1, 5),
                ]
            )

    story.extend(
        [
            Spacer(1, 10),
            Paragraph(
                "Generated by EduAI Assignment Evaluation",
                small_style,
            ),
        ]
    )

    document.build(
        story
    )

    return buffer.getvalue()


def build_evaluation_report_download(
    evaluation,
    snapshot,
    report_format,
):
    """Return bytes, filename and MIME type for the selected report format."""

    normalized_format = str(
        report_format
    ).strip().upper()

    if normalized_format == "PDF":
        return (
            build_evaluation_report_pdf(
                evaluation,
                snapshot,
            ),
            evaluation_report_filename(
                snapshot,
                "pdf",
            ),
            "application/pdf",
        )

    if normalized_format in {
        "WORD (DOCX)",
        "DOCX",
        "WORD",
    }:
        return (
            build_evaluation_report_docx(
                evaluation,
                snapshot,
            ),
            evaluation_report_filename(
                snapshot,
                "docx",
            ),
            (
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

    report_text = build_evaluation_report_text(
        evaluation,
        snapshot,
    )

    return (
        report_text.encode(
            "utf-8"
        ),
        evaluation_report_filename(
            snapshot,
            "txt",
        ),
        "text/plain",
    )




# =========================================================
# QUESTION BANK HELPERS
# =========================================================

def _download_response_bytes(download_data):
    """
    Normalize a Question Bank PDF download response to raw bytes.
    """

    if download_data is None:
        return None

    if isinstance(
        download_data,
        bytes,
    ):
        return download_data

    if isinstance(
        download_data,
        bytearray,
    ):
        return bytes(
            download_data
        )

    if hasattr(
        download_data,
        "content",
    ):

        content = download_data.content

        if isinstance(
            content,
            bytes,
        ):
            return content

        if isinstance(
            content,
            bytearray,
        ):
            return bytes(
                content
            )

    if hasattr(
        download_data,
        "getvalue",
    ):
        return download_data.getvalue()

    if hasattr(
        download_data,
        "read",
    ):

        try:
            download_data.seek(
                0
            )
        except Exception:
            pass

        return download_data.read()

    return None


def _normalize_question_bank_set(
    data,
    fallback_id=None,
):
    """
    Normalize Question Bank detail responses.
    """

    if not isinstance(
        data,
        dict,
    ):
        return {}

    if isinstance(
        data.get(
            "question_set"
        ),
        dict,
    ):

        normalized = data[
            "question_set"
        ].copy()

        if isinstance(
            data.get(
                "questions"
            ),
            list,
        ):

            normalized[
                "questions"
            ] = data[
                "questions"
            ]

    else:

        normalized = data.copy()

    existing_id = (
        normalized.get(
            "id"
        )
        or
        normalized.get(
            "set_id"
        )
        or
        fallback_id
    )

    if existing_id is not None:
        normalized[
            "id"
        ] = existing_id

    return normalized


def _question_bank_assignment_text(
    selected_set,
):
    """
    Reconstruct a saved Question Bank assignment as text if its
    PDF export endpoint is unavailable.
    """

    title = str(
        selected_set.get(
            "title",
            "EduAI Assignment",
        )
        or
        "EduAI Assignment"
    )

    topic = str(
        selected_set.get(
            "topic",
            "-",
        )
        or
        "-"
    )

    difficulty = str(
        selected_set.get(
            "difficulty",
            "-",
        )
        or
        "-"
    )

    questions = selected_set.get(
        "questions",
        [],
    )

    if not isinstance(
        questions,
        list,
    ):
        questions = []

    total_marks = selected_set.get(
        "total_marks",
        0,
    )

    try:
        numeric_total = float(
            total_marks
            or
            0
        )
    except (
        TypeError,
        ValueError,
    ):
        numeric_total = 0.0

    if (
        numeric_total <= 0
        and
        questions
    ):

        calculated_total = 0.0

        for question in questions:

            try:
                calculated_total += float(
                    question.get(
                        "marks",
                        0,
                    )
                    or
                    0
                )
            except (
                TypeError,
                ValueError,
            ):
                continue

        numeric_total = calculated_total

    if float(
        numeric_total
    ).is_integer():

        total_marks_text = str(
            int(
                numeric_total
            )
        )

    else:

        total_marks_text = str(
            numeric_total
        )

    lines = [
        title,
        "=" * max(
            20,
            len(
                title
            ),
        ),
        "",
        f"Topic: {topic}",
        f"Difficulty: {difficulty}",
        f"Total Marks: {total_marks_text}",
        "",
        "Questions",
        "-" * 40,
        "",
    ]

    for index, question in enumerate(
        questions,
        start=1,
    ):

        number = question.get(
            "question_number",
            index,
        )

        question_text = (
            question.get(
                "question_text"
            )
            or
            question.get(
                "question"
            )
            or
            ""
        )

        question_type = (
            question.get(
                "type"
            )
            or
            "-"
        )

        marks = question.get(
            "marks",
            0,
        )

        lines.append(
            f"Question {number}: {question_text}"
        )

        options = question.get(
            "options"
        )

        if isinstance(
            options,
            str,
        ):

            try:

                parsed_options = json.loads(
                    options
                )

                if isinstance(
                    parsed_options,
                    list,
                ):
                    options = parsed_options

            except Exception:

                options = [
                    options
                ]

        if isinstance(
            options,
            (
                list,
                tuple,
            ),
        ):

            for option in options:

                lines.append(
                    f"  - {option}"
                )

        lines.append(
            f"Type: {question_type} | Marks: {marks}"
        )

        lines.append(
            ""
        )

    return "\n".join(
        lines
    )


def build_question_bank_assignment_file(
    question_set_id,
    question_set_summary=None,
):
    """
    Convert a saved Question Bank record into the same file-like
    assignment object used by the existing upload flow.

    The saved PDF is preferred. If PDF export fails, a clean TXT
    representation is created from the saved question data.
    """

    details = get_question_set(
        question_set_id
    )

    selected_set = _normalize_question_bank_set(
        details,
        fallback_id=question_set_id,
    )

    if (
        not selected_set
        and
        isinstance(
            question_set_summary,
            dict,
        )
    ):

        selected_set = question_set_summary.copy()

    title = str(
        selected_set.get(
            "title",
            f"Question_Set_{question_set_id}",
        )
        or
        f"Question_Set_{question_set_id}"
    )

    safe_title = re.sub(
        r"[^A-Za-z0-9_-]+",
        "_",
        title,
    ).strip(
        "_"
    )

    if not safe_title:
        safe_title = (
            f"Question_Set_{question_set_id}"
        )

    # -----------------------------------------------------
    # Prefer Question Bank PDF export.
    # -----------------------------------------------------

    try:

        pdf_data = _download_response_bytes(
            download_pdf(
                question_set_id
            )
        )

        if (
            pdf_data
            and
            pdf_data.startswith(
                b"%PDF"
            )
        ):

            assignment_file = io.BytesIO(
                pdf_data
            )

            assignment_file.name = (
                f"{safe_title}.pdf"
            )

            assignment_file.type = (
                "application/pdf"
            )

            return (
                assignment_file,
                selected_set,
                "pdf",
            )

    except Exception as pdf_error:

        print(
            "Question Bank PDF export fallback:",
            pdf_error,
        )

    # -----------------------------------------------------
    # Safe fallback from saved question data.
    # -----------------------------------------------------

    assignment_text = (
        _question_bank_assignment_text(
            selected_set
        )
    )

    assignment_file = io.BytesIO(
        assignment_text.encode(
            "utf-8"
        )
    )

    assignment_file.name = (
        f"{safe_title}.txt"
    )

    assignment_file.type = (
        "text/plain"
    )

    return (
        assignment_file,
        selected_set,
        "txt",
    )


def question_bank_option_label(
    item,
):
    """
    Readable selectbox label for a saved Question Bank assignment.
    """

    item_id = (
        item.get(
            "id"
        )
        or
        item.get(
            "set_id"
        )
        or
        "-"
    )

    title = str(
        item.get(
            "title",
            "Untitled Question Set",
        )
    )

    topic = str(
        item.get(
            "topic",
            "-",
        )
    )

    total_marks = item.get(
        "total_marks",
        "-"
    )

    return (
        f"{title}  |  {topic}  |  "
        f"{total_marks} marks  |  ID {item_id}"
    )


def create_file_snapshot(uploaded_file):
    if uploaded_file is None:
        return None

    try:
        file_bytes = uploaded_file.getvalue()
    except Exception:
        try:
            uploaded_file.seek(0)
        except Exception:
            pass
        file_bytes = uploaded_file.read()

    return {
        "bytes": file_bytes,
        "name": getattr(uploaded_file, "name", "uploaded_file"),
        "type": getattr(
            uploaded_file,
            "type",
            "application/octet-stream",
        ),
    }


def rebuild_uploaded_file(file_snapshot):
    if not file_snapshot:
        return None

    rebuilt = io.BytesIO(file_snapshot["bytes"])
    rebuilt.name = file_snapshot.get("name", "uploaded_file")
    rebuilt.type = file_snapshot.get(
        "type",
        "application/octet-stream",
    )
    return rebuilt




# =========================================================
# STUDENT SUBMISSION HELPERS
# =========================================================

def _uploaded_file_bytes(uploaded_file):
    """
    Read a Streamlit uploaded file without depending on its current cursor.
    """

    try:
        return uploaded_file.getvalue()
    except Exception:
        try:
            uploaded_file.seek(0)
        except Exception:
            pass
        return uploaded_file.read()


def _merge_pdf_byte_streams(pdf_items):
    """
    Merge PDF byte streams into one PDF while preserving all pages.
    """

    try:
        import pymupdf
    except ImportError:
        try:
            import fitz as pymupdf
        except ImportError as error:
            raise RuntimeError(
                "Multiple-PDF and ZIP submissions require PyMuPDF. "
                "Install it with: python -m pip install pymupdf"
            ) from error

    output_document = pymupdf.open()

    try:
        for file_name, pdf_bytes in pdf_items:

            if not pdf_bytes:
                raise ValueError(
                    f"{file_name} is empty."
                )

            try:
                source_document = pymupdf.open(
                    stream=pdf_bytes,
                    filetype="pdf",
                )
            except Exception as error:
                raise ValueError(
                    f"{file_name} is not a readable PDF."
                ) from error

            try:
                if source_document.page_count <= 0:
                    raise ValueError(
                        f"{file_name} does not contain any PDF pages."
                    )

                output_document.insert_pdf(
                    source_document
                )

            finally:
                source_document.close()

        if output_document.page_count <= 0:
            raise ValueError(
                "No readable PDF pages were found in the submission."
            )

        return output_document.tobytes(
            garbage=4,
            deflate=True,
        )

    finally:
        output_document.close()


def validate_submission_selection(uploaded_files):
    """
    Allowed Student Submission combinations:
      - one PDF
      - one DOCX
      - multiple PDFs/DOCX files
      - one ZIP containing supported files
    """

    files = list(
        uploaded_files
        or []
    )

    if not files:
        return {
            "valid": False,
            "message": None,
            "mode": None,
        }

    extensions = [
        os.path.splitext(
            getattr(
                file,
                "name",
                "",
            )
        )[1].lower()
        for file in files
    ]

    if any(
        extension not in {
            ".pdf",
            ".docx",
            ".zip",
        }
        for extension in extensions
    ):
        return {
            "valid": False,
            "message": (
                "Student Submission accepts PDF, DOCX files or one ZIP file only."
            ),
            "mode": None,
        }

    zip_count = extensions.count(
        ".zip"
    )

    if zip_count:

        if len(files) > 1:
            return {
                "valid": False,
                "message": (
                    "Upload either multiple PDF files or one ZIP file. "
                    "Do not mix a ZIP file with separate PDFs."
                ),
                "mode": None,
            }

        return {
            "valid": True,
            "message": None,
            "mode": "zip",
            "count": 1,
        }

    return {
        "valid": True,
        "message": None,
        "mode": "pdf",
        "count": len(files),
    }


def prepare_submission_snapshots(uploaded_files):
    """
    Prepare EACH student PDF as an independent submission.

    - One PDF -> one student evaluation.
    - Multiple PDFs -> one evaluation per PDF.
    - One ZIP -> every PDF inside the ZIP becomes one independent
      student evaluation.

    Files are NOT merged. The active assignment and rubric are
    applied separately to every student submission.
    """

    files = list(
        uploaded_files
        or []
    )

    validation = validate_submission_selection(
        files
    )

    if not validation.get(
        "valid"
    ):
        raise ValueError(
            validation.get(
                "message"
            )
            or
            "Please upload a valid student submission."
        )

    prepared = []

    # -----------------------------------------------------
    # ZIP -> MANY INDEPENDENT PDF SUBMISSIONS
    # -----------------------------------------------------

    if validation.get(
        "mode"
    ) == "zip":

        zip_file = files[0]
        zip_bytes = _uploaded_file_bytes(
            zip_file
        )

        try:

            with zipfile.ZipFile(
                io.BytesIO(
                    zip_bytes
                )
            ) as archive:

                for info in archive.infolist():

                    if info.is_dir():
                        continue

                    if not info.filename.lower().endswith(
                        ".pdf"
                    ):
                        continue

                    if info.file_size > 100 * 1024 * 1024:
                        raise ValueError(
                            f"{info.filename} is larger than "
                            "the 100 MB per-file submission limit."
                        )

                    file_bytes = archive.read(
                        info
                    )

                    if not file_bytes:
                        continue

                    prepared.append(
                        {
                            "bytes": file_bytes,
                            "name": os.path.basename(
                                info.filename
                            ),
                            "type": "application/pdf",
                            "source_mode": "zip",
                            "original_zip_name": getattr(
                                zip_file,
                                "name",
                                "submission.zip",
                            ),
                        }
                    )

        except zipfile.BadZipFile as error:

            raise ValueError(
                "The selected ZIP file is invalid or corrupted."
            ) from error

        if not prepared:

            raise ValueError(
                "The ZIP file does not contain any PDF submissions."
            )

        return prepared

    # -----------------------------------------------------
    # ONE OR MULTIPLE PDFs -> ONE SNAPSHOT PER PDF
    # -----------------------------------------------------

    for uploaded_file in files:

        snapshot = create_file_snapshot(
            uploaded_file
        )

        snapshot[
            "source_mode"
        ] = (
            "single_pdf"
            if len(files) == 1
            else
            "multiple_pdf"
        )

        prepared.append(
            snapshot
        )

    return prepared



def _excel_column_name(column_number):
    # Convert a 1-based column number to an Excel column label.
    result = ""

    while column_number:
        column_number, remainder = divmod(
            column_number - 1,
            26,
        )
        result = chr(65 + remainder) + result

    return result


def _xlsx_inline_cell(
    cell_reference,
    value,
    style_id=0,
):
    # Build a safe text cell for the XLSX XML.
    escaped_value = html.escape(
        "" if value is None else str(value),
        quote=False,
    )

    return (
        f'<c r="{cell_reference}" t="inlineStr" s="{style_id}">'
        f'<is><t xml:space="preserve">{escaped_value}</t></is></c>'
    )


def _xlsx_number_cell(
    cell_reference,
    value,
    style_id=0,
):
    # Build a numeric XLSX cell.
    if value is None:
        return (
            f'<c r="{cell_reference}" '
            f't="inlineStr" s="{style_id}"><is><t></t></is></c>'
        )

    numeric_value = safe_number(
        value,
        0,
    )

    return (
        f'<c r="{cell_reference}" s="{style_id}">'
        f'<v>{numeric_value}</v></c>'
    )


def _student_name_for_excel(
    evaluation,
    submission_snapshot,
):
    # Prefer an explicit student_name; otherwise derive it from the PDF name.
    explicit_name = str(
        (evaluation or {}).get(
            "student_name",
            "",
        )
        or ""
    ).strip()

    if explicit_name:
        return explicit_name

    submission_name = str(
        (submission_snapshot or {}).get(
            "name",
            "Student",
        )
        or
        "Student"
    )

    base_name = os.path.splitext(
        os.path.basename(
            submission_name
        )
    )[0]

    readable_name = (
        base_name
        .replace("_", " ")
        .replace("-", " ")
    )

    return " ".join(
        readable_name.split()
    ) or "Student"


def _safe_identifier(value, fallback):
    text = re.sub(r"[^A-Za-z0-9_-]+", "_", str(value or "").strip()).strip("_")
    return text or fallback


def build_report_ready_batch_rows(batch_results, assignment_snapshot):
    """Build the exact row schema required by Report Generation."""
    rows = []
    assignment_filename = (assignment_snapshot or {}).get("name", "Assignment")
    assignment_stem = os.path.splitext(os.path.basename(str(assignment_filename)))[0]
    evaluation_date = datetime.now().strftime("%Y-%m-%d")

    for index, item in enumerate(batch_results or [], start=1):
        evaluation = item.get("evaluation")
        if not evaluation:
            continue
        submission_snapshot = item.get("submission_snapshot", {})
        submission_name = submission_snapshot.get("name", f"student_{index}.pdf")
        student_name = _student_name_for_excel(evaluation, submission_snapshot)
        student_id = _safe_identifier(
            evaluation.get("student_id") or os.path.splitext(os.path.basename(submission_name))[0],
            f"STUDENT_{index}",
        )
        course_id = _safe_identifier(evaluation.get("course_id"), "COURSE_UNASSIGNED")
        cohort_id = _safe_identifier(evaluation.get("cohort_id") or evaluation.get("cohort"), "COHORT_UNASSIGNED")
        assignment_id = _safe_identifier(
            evaluation.get("assignment_id") or evaluation.get("assignment_name") or assignment_stem,
            "ASSIGNMENT",
        )
        remarks = evaluation.get("overall_feedback") or evaluation.get("remarks") or "No remarks provided."
        criteria = evaluation.get("criteria") or []

        if criteria:
            for c_index, criterion in enumerate(criteria, start=1):
                rows.append({
                    "student_id": student_id,
                    "student_name": student_name,
                    "course_id": course_id,
                    "cohort_id": cohort_id,
                    "assignment_id": assignment_id,
                    "criteria": criterion.get("criterion") or criterion.get("name") or f"Criterion {c_index}",
                    "max_marks": safe_number(criterion.get("max_score", criterion.get("max_marks", 0))),
                    "obtained_marks": safe_number(criterion.get("score", criterion.get("obtained_marks", 0))),
                    "evaluation_date": evaluation_date,
                    "remarks": criterion.get("feedback") or remarks,
                })
        else:
            rows.append({
                "student_id": student_id,
                "student_name": student_name,
                "course_id": course_id,
                "cohort_id": cohort_id,
                "assignment_id": assignment_id,
                "criteria": "Overall",
                "max_marks": safe_number(evaluation.get("max_score", evaluation.get("total_marks", 0))),
                "obtained_marks": safe_number(evaluation.get("total_score", evaluation.get("obtained_marks", 0))),
                "evaluation_date": evaluation_date,
                "remarks": remarks,
            })

    return rows


def build_batch_marks_xlsx(batch_results, assignment_snapshot):
    """Export the exact mandatory Report Generation schema to XLSX."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment

    rows = build_report_ready_batch_rows(batch_results, assignment_snapshot)
    if not rows:
        raise RuntimeError("No successful batch evaluation data is available for XLSX export.")

    headers = [
        "student_id", "student_name", "course_id", "cohort_id",
        "assignment_id", "criteria", "max_marks", "obtained_marks",
        "evaluation_date", "remarks",
    ]
    wb = Workbook()
    ws = wb.active
    ws.title = "Evaluation Results"
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal="center")
    for row in rows:
        ws.append([row.get(h) for h in headers])
    ws.freeze_panes = "A2"
    widths = {
        "A":18,"B":24,"C":20,"D":20,"E":26,"F":28,
        "G":14,"H":16,"I":18,"J":55,
    }
    for col, width in widths.items():
        ws.column_dimensions[col].width = width
    for row in ws.iter_rows(min_row=2):
        row[9].alignment = Alignment(wrap_text=True, vertical="top")

    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


def batch_marks_xlsx_filename(
    batch_results,
):
    # Use the original uploaded ZIP name when available.

    for item in (
        batch_results
        or []
    ):

        submission_snapshot = item.get(
            "submission_snapshot",
            {},
        )

        original_zip_name = submission_snapshot.get(
            "original_zip_name"
        )

        if original_zip_name:

            zip_base = os.path.splitext(
                os.path.basename(
                    original_zip_name
                )
            )[0]

            safe_base = re.sub(
                r"[^A-Za-z0-9_-]+",
                "_",
                zip_base,
            ).strip(
                "_"
            )

            if safe_base:

                return (
                    f"{safe_base}_Evaluation_Summary.xlsx"
                )

    return "EduAI_Batch_Evaluation_Summary.xlsx"


def build_batch_reports_zip(
    batch_results,
    assignment_snapshot,
    rubric_snapshot,
    report_format,
):
    """
    Build one ZIP containing one evaluation report per successful student.
    """

    output = io.BytesIO()
    successful_count = 0
    failed_lines = []

    with zipfile.ZipFile(
        output,
        mode="w",
        compression=zipfile.ZIP_DEFLATED,
    ) as archive:

        for item in batch_results:

            submission_snapshot = item.get(
                "submission_snapshot"
            )

            evaluation = item.get(
                "evaluation"
            )

            error_message = item.get(
                "error"
            )

            submission_name = (
                (submission_snapshot or {})
                .get(
                    "name",
                    "student_submission.pdf",
                )
            )

            if not evaluation:

                failed_lines.append(
                    f"{submission_name}: "
                    f"{error_message or 'Evaluation failed.'}"
                )

                continue

            per_student_snapshot = {
                "assignment_file": assignment_snapshot,
                "rubric_file": rubric_snapshot,
                "submission_file": submission_snapshot,
            }

            (
                report_bytes,
                report_file_name,
                _,
            ) = build_evaluation_report_download(
                evaluation=evaluation,
                snapshot=per_student_snapshot,
                report_format=report_format,
            )

            archive.writestr(
                report_file_name,
                report_bytes,
            )

            successful_count += 1

        if failed_lines:

            archive.writestr(
                "Failed_Evaluations.txt",
                (
                    "EduAI - Failed Batch Evaluations\n"
                    "================================\n\n"
                    + "\n".join(
                        failed_lines
                    )
                ).encode(
                    "utf-8"
                ),
            )

    if successful_count == 0:

        raise RuntimeError(
            "No successful evaluation reports are available to download."
        )

    output.seek(
        0
    )

    return output.getvalue()



def submission_selection_label(uploaded_files):
    """
    Human-friendly Student Submission selection text.
    """

    files = list(
        uploaded_files
        or []
    )

    if not files:
        return "Submission not selected"

    validation = validate_submission_selection(
        files
    )

    if not validation.get(
        "valid"
    ):
        return (
            validation.get(
                "message"
            )
            or
            "Invalid submission selection"
        )

    if validation.get(
        "mode"
    ) == "zip":
        return (
            "Selected batch ZIP: "
            + getattr(
                files[0],
                "name",
                "submission.zip",
            )
        )

    if len(files) == 1:
        return (
            "Selected: "
            + getattr(
                files[0],
                "name",
                "submission.pdf",
            )
        )

    return (
        f"Selected batch: {len(files)} student PDF files"
    )


# =========================================================
# RUBRIC HELPERS
# =========================================================

def calculate_file_hash(file_snapshot):
    if not file_snapshot:
        return None

    return hashlib.sha256(
        file_snapshot["bytes"]
    ).hexdigest()


def generate_rubric_from_assignment(file_snapshot):
    """
    Send the uploaded assignment through api_client.py to:
        POST /assignment-evaluation/generate-rubric

    The backend then extracts the assignment text and calls the
    configured Gemini rubric-generation service.
    """

    if not file_snapshot:
        raise ValueError(
            "Assignment file is required for rubric generation."
        )

    assignment_file = rebuild_uploaded_file(
        file_snapshot
    )

    return generate_assignment_rubric(
        assignment_file=assignment_file
    )


def normalize_rubric(api_result):
    if not isinstance(api_result, dict):
        raise ValueError("Rubric Agent returned an invalid response.")

    source = api_result.get("rubric", api_result)
    if not isinstance(source, dict):
        raise ValueError("Rubric Agent response does not contain a rubric object.")

    title = (
        source.get("title")
        or source.get("rubric_title")
        or source.get("assignment_title")
        or "AI Generated Rubric"
    )
    total_marks = safe_number(
        source.get("total_marks", source.get("total_points", 0))
    )
    raw_criteria = source.get("criteria") or source.get("rubric_criteria") or []
    if not isinstance(raw_criteria, list):
        raw_criteria = []

    criteria=[]
    for index, criterion in enumerate(raw_criteria, start=1):
        if not isinstance(criterion, dict):
            continue
        name=(criterion.get("name") or criterion.get("title") or criterion.get("criterion") or f"Criterion {index}")
        description=(criterion.get("description") or criterion.get("objective") or "")
        weight=safe_number(criterion.get("weight", criterion.get("weight_percent", criterion.get("weightage", 0))))
        max_marks=safe_number(criterion.get("max_marks", criterion.get("marks", 0)))
        if max_marks <= 0 and total_marks > 0 and weight > 0:
            max_marks=total_marks*weight/100
        levels=criterion.get("levels", criterion.get("performance_levels", []))
        criteria.append({
            "name": str(name),
            "description": str(description),
            "max_marks": max_marks,
            "weight": weight,
            "levels": levels if isinstance(levels, list) else [],
        })

    if total_marks <= 0 and criteria:
        total_marks=sum(safe_number(item.get("max_marks",0)) for item in criteria)
    if total_marks <= 0:
        raise ValueError("The generated rubric does not contain valid total marks.")
    if not criteria:
        raise ValueError("The generated rubric does not contain any criteria.")

    for item in criteria:
        if safe_number(item.get("max_marks")) <= 0 and safe_number(item.get("weight")) > 0:
            item["max_marks"]=total_marks*safe_number(item["weight"])/100

    marks_sum=sum(safe_number(item.get("max_marks",0)) for item in criteria)
    if marks_sum > 0 and abs(marks_sum-total_marks) > 0.05:
        scale=total_marks/marks_sum
        running=0.0
        for idx,item in enumerate(criteria):
            if idx == len(criteria)-1:
                item["max_marks"]=total_marks-running
            else:
                adjusted=round(safe_number(item["max_marks"])*scale,2)
                item["max_marks"]=adjusted
                running += adjusted

    for item in criteria:
        item["weight"]=(safe_number(item.get("max_marks",0))/total_marks*100) if total_marks else 0

    return {"title":str(title),"total_marks":total_marks,"criteria":criteria}


def validate_rubric(rubric):
    if not isinstance(rubric, dict):
        return False, "Rubric is invalid."
    criteria=rubric.get("criteria",[])
    if not criteria:
        return False, "Rubric must contain at least one criterion."
    total=safe_number(rubric.get("total_marks",0))
    if total <= 0:
        return False, "Rubric total marks must be greater than zero."
    marks_sum=sum(safe_number(item.get("max_marks",0)) for item in criteria)
    if abs(marks_sum-total) > 0.05:
        return False, f"Criterion marks ({format_score(marks_sum)}) must equal rubric total ({format_score(total)})."
    for i,item in enumerate(criteria,start=1):
        if not str(item.get("name","")).strip():
            return False, f"Criterion {i} must have a name."
    return True, ""


def rubric_to_text(rubric):
    lines=[
        f"RUBRIC TITLE: {rubric.get('title','AI Generated Rubric')}",
        f"TOTAL MARKS: {format_score(rubric.get('total_marks',0))}",
        "",
        "CRITERIA",
        "",
    ]
    for i,item in enumerate(rubric.get("criteria",[]),start=1):
        lines.extend([
            f"{i}. {item.get('name',f'Criterion {i}')}",
            f"Description: {item.get('description','')}",
            f"Maximum Marks: {format_score(item.get('max_marks',0))}",
            f"Weight: {safe_number(item.get('weight',0)):.2f}%",
            "",
        ])
    return "\n".join(lines)


def create_generated_rubric_snapshot(rubric):
    text=rubric_to_text(rubric)
    return {"bytes":text.encode("utf-8"),"name":"AI_Generated_Rubric.txt","type":"text/plain"}


def get_active_rubric_snapshot():
    if st.session_state.get("rubric_source") == "uploaded":
        return st.session_state.get("uploaded_rubric_snapshot")
    rubric=st.session_state.get("generated_rubric")
    return create_generated_rubric_snapshot(rubric) if rubric else None


def reset_rubric_flow():
    st.session_state.generated_rubric = None
    st.session_state.rubric_generation_phase = "idle"
    st.session_state.rubric_generation_error = None
    st.session_state.rubric_source = None
    st.session_state.uploaded_rubric_snapshot = None
    st.session_state.show_rubric_upload = False

    # Remove stale edit-widget values when a different
    # assignment generates a different rubric.
    edit_keys = [
        key
        for key in st.session_state.keys()
        if (
            key.startswith("rubric_edit_")
            or key in {
                "edit_rubric_title",
                "edit_rubric_total_marks",
            }
        )
    ]

    for key in edit_keys:
        del st.session_state[key]


# =========================================================
# CONSISTENT EDUAI UI
# =========================================================

st.html(
    """
<style>

[data-testid="stSidebar"] {
    display: none !important;
}

[data-testid="collapsedControl"] {
    display: none !important;
}

header {
    display: none !important;
}

footer {
    display: none !important;
}

#MainMenu {
    visibility: hidden !important;
}

.stApp {
    background:
        radial-gradient(
            circle at 8% 12%,
            rgba(99, 102, 241, 0.08),
            transparent 28%
        ),
        radial-gradient(
            circle at 92% 82%,
            rgba(249, 115, 22, 0.06),
            transparent 30%
        ),
        #f8fafc;
}

.block-container {
    max-width: 1180px !important;
    padding-top: 45px !important;
    padding-bottom: 80px !important;
    padding-left: 28px !important;
    padding-right: 28px !important;
}



/* =====================================================
   VIEW LOGS BUTTON
===================================================== */

.st-key-ai_eval_view_logs button {

    width: 150px !important;
    min-height: 40px !important;
    height: 40px !important;

    background: #ffffff !important;
    color: #4f46e5 !important;

    border: 1px solid #c7d2fe !important;
    border-radius: 10px !important;

    font-size: 12.5px !important;
    font-weight: 700 !important;

    box-shadow:
        0 3px 10px
        rgba(15, 23, 42, 0.05) !important;

    transition: all 0.18s ease !important;
}


.st-key-ai_eval_view_logs button:hover {

    background: #4f46e5 !important;
    color: #ffffff !important;

    border-color: #4f46e5 !important;

    transform: translateY(-1px);

    box-shadow:
        0 7px 16px
        rgba(79, 70, 229, 0.16) !important;
}


/* =====================================================
   EVALUATION LOGS
===================================================== */

.eval-logs-header {
    margin-bottom: 14px;
}

.eval-logs-title {
    color: #0f172a;
    font-size: 22px;
    font-weight: 780;
    margin-bottom: 6px;
}

.eval-logs-description {
    color: #64748b;
    font-size: 13px;
    line-height: 1.6;
}

.eval-log-summary {
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 13px;
    padding: 13px 15px;
    margin: 10px 0 16px 0;
}

.eval-log-card {
    background: #ffffff;
    border: 1px solid #e2e8f0;
    border-radius: 14px;
    padding: 16px 17px;
    margin-bottom: 10px;
}

.eval-log-title {
    color: #0f172a;
    font-size: 14px;
    font-weight: 750;
    margin-bottom: 5px;
}

.eval-log-meta {
    color: #64748b;
    font-size: 12px;
    line-height: 1.7;
}

.eval-log-score {
    color: #4f46e5;
    font-size: 12.5px;
    font-weight: 700;
    margin-top: 7px;
}

.st-key-ai_eval_logs_scroll {
    height: 64vh !important;
    max-height: 64vh !important;
    overflow-y: auto !important;
    overflow-x: hidden !important;
    padding-right: 8px !important;
    scrollbar-width: thin;
    scrollbar-color: #cbd5e1 transparent;
}

/* =====================================================
   BACK TO DASHBOARD
===================================================== */

.back-link-wrap {
    padding-top: 28px;
    display: flex;
    justify-content: flex-start;
}

.back-link {
    width: 40px;
    height: 40px;

    display: inline-flex;
    align-items: center;
    justify-content: center;

    background: #ffffff;
    color: #334155 !important;

    border: 1px solid #dbe3ed;
    border-radius: 12px;

    font-size: 24px;
    font-weight: 500;
    line-height: 1;

    text-decoration: none !important;

    box-shadow:
        0 4px 12px
        rgba(15, 23, 42, 0.06);

    transition:
        all 0.18s ease;
}

.back-link:hover {
    background: #f8fafc;
    color: #4f46e5 !important;
    border-color: #c7d2fe;

    transform: translateX(-1px);

    box-shadow:
        0 6px 16px
        rgba(15, 23, 42, 0.08);
}

.back-link:focus {
    outline: none;

    box-shadow:
        0 0 0 3px
        rgba(79, 70, 229, 0.10);
}


.page-header {
    margin-bottom: 30px;
}

.page-label {
    color: #4f46e5;
    font-size: 12px;
    font-weight: 750;
    letter-spacing: 1.2px;
    text-transform: uppercase;
    margin-bottom: 12px;
}

.page-title {
    color: #0f172a;
    font-size: 38px;
    font-weight: 780;
    letter-spacing: -1.1px;
    line-height: 1.15;
    margin-bottom: 12px;
}

.page-description {
    color: #64748b;
    font-size: 15px;
    line-height: 1.7;
    max-width: 780px;
}

.status-row {
    display: flex;
    align-items: center;
    gap: 8px;
    margin-top: 15px;
    font-size: 13px;
    color: #64748b;
}

.status-dot {
    width: 8px;
    height: 8px;
    border-radius: 50%;
    background: #16a34a;
    box-shadow: 0 0 0 4px rgba(22, 163, 74, 0.10);
}

.status-dot.offline {
    background: #dc2626;
    box-shadow: 0 0 0 4px rgba(220, 38, 38, 0.09);
}

[data-testid="stVerticalBlockBorderWrapper"] {
    background: rgba(255, 255, 255, 0.97);
    border: 1px solid #e2e8f0 !important;
    border-radius: 20px !important;
    box-shadow: 0 8px 28px rgba(15, 23, 42, 0.045);
    margin-bottom: 20px;
}

.section-header {
    padding: 3px 1px 16px 1px;
}

.section-number {
    color: #4f46e5;
    font-size: 11px;
    font-weight: 750;
    letter-spacing: 1px;
    text-transform: uppercase;
    margin-bottom: 8px;
}

.section-title {
    color: #0f172a;
    font-size: 19px;
    font-weight: 760;
    letter-spacing: -0.3px;
    margin-bottom: 6px;
}

.section-description {
    color: #64748b;
    font-size: 13.5px;
    line-height: 1.6;
}

.upload-label {
    color: #0f172a;
    font-size: 14px;
    font-weight: 700;
    margin-bottom: 5px;
}

.upload-help {
    color: #64748b;
    font-size: 12.5px;
    min-height: 42px;
    line-height: 1.5;
}

[data-testid="stFileUploaderDropzone"] {
    background: #f8fafc !important;
    border: 1px dashed #cbd5e1 !important;
    border-radius: 14px !important;
}

[data-testid="stFileUploaderDropzone"]:hover {
    border-color: #818cf8 !important;
    background: #f5f7ff !important;
}

.stButton > button {
    min-height: 46px !important;
    border-radius: 12px !important;
    border: 1px solid #4f46e5 !important;
    background: #4f46e5 !important;
    color: #ffffff !important;
    font-size: 14px !important;
    font-weight: 700 !important;
    box-shadow: 0 8px 20px rgba(79, 70, 229, 0.16) !important;
    transition: all 0.2s ease !important;
}

.stButton > button:hover {
    background: #4338ca !important;
    border-color: #4338ca !important;
    transform: translateY(-1px);
}

.file-pill {
    border: 1px solid #e2e8f0;
    background: #f8fafc;
    border-radius: 12px;
    padding: 12px 14px;
    color: #475569;
    font-size: 13px;
    min-height: 46px;
}

.file-pill.ready {
    background: #f0fdf4;
    border-color: #bbf7d0;
    color: #166534;
}

.question-bank-card {
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 13px;
    padding: 14px 16px;
    margin-top: 10px;
}

.question-bank-title {
    color: #0f172a;
    font-size: 13.5px;
    font-weight: 740;
    margin-bottom: 5px;
}

.question-bank-meta {
    color: #64748b;
    font-size: 12.5px;
    line-height: 1.6;
}

.result-banner {
    background: linear-gradient(135deg, #eef2ff, #ffffff);
    border: 1px solid #c7d2fe;
    border-left: 4px solid #4f46e5;
    border-radius: 16px;
    padding: 19px 20px;
    margin-bottom: 20px;
}

.result-title {
    color: #0f172a;
    font-size: 20px;
    font-weight: 760;
    margin-bottom: 5px;
}

.result-text {
    color: #64748b;
    font-size: 13.5px;
}

[data-testid="stMetric"] {
    background: #ffffff;
    border: 1px solid #e2e8f0;
    border-radius: 16px;
    padding: 17px 18px;
    box-shadow: 0 5px 18px rgba(15, 23, 42, 0.035);
}

[data-testid="stMetricLabel"] {
    color: #64748b !important;
}

[data-testid="stMetricValue"] {
    color: #0f172a !important;
    font-weight: 780 !important;
}

.remark-card {
    background: #ffffff;
    border: 1px solid #e2e8f0;
    border-radius: 15px;
    padding: 18px 20px;
    color: #475569;
    font-size: 13.5px;
    line-height: 1.7;
}

.remark-label {
    color: #4f46e5;
    font-size: 11px;
    font-weight: 750;
    letter-spacing: 0.9px;
    text-transform: uppercase;
    margin-bottom: 7px;
}

.deduction-card {
    background: #ffffff;
    border: 1px solid #e2e8f0;
    border-left: 4px solid #f59e0b;
    border-radius: 14px;
    padding: 16px 18px;
    margin-bottom: 10px;
}

.deduction-question {
    color: #0f172a;
    font-size: 14px;
    font-weight: 730;
}

.deduction-marks {
    color: #dc2626;
    font-size: 12.5px;
    font-weight: 700;
    margin-top: 5px;
}

.deduction-reason {
    color: #64748b;
    font-size: 13px;
    line-height: 1.6;
    margin-top: 7px;
}



/* =====================================================
   ASSIGNMENT-DRIVEN RUBRIC FLOW
===================================================== */

.rubric-ready-card {
    background: linear-gradient(135deg, #ffffff, #f8faff);
    border: 1px solid #c7d2fe;
    border-radius: 16px;
    padding: 18px 20px;
    margin: 8px 0 14px 0;
}

.rubric-ready-title {
    color: #0f172a;
    font-size: 16px;
    font-weight: 760;
    margin-bottom: 6px;
}

.rubric-ready-meta {
    color: #64748b;
    font-size: 12.5px;
    line-height: 1.65;
}

.rubric-source-badge {
    display: inline-flex;
    align-items: center;
    padding: 5px 9px;
    border-radius: 999px;
    background: #eef2ff;
    color: #4f46e5;
    border: 1px solid #c7d2fe;
    font-size: 11px;
    font-weight: 750;
    margin-bottom: 9px;
}

.rubric-total-bar {
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 12px;
    padding: 12px 14px;
    color: #475569;
    font-size: 12.5px;
    margin: 12px 0;
}

/* Rubric generation modal preview */
.rubric-preview-success {
    background: #f0fdf4;
    border: 1px solid #bbf7d0;
    border-radius: 13px;
    padding: 12px 15px;
    color: #166534;
    font-size: 13px;
    font-weight: 650;
    margin: 8px 0 14px 0;
}

.rubric-preview-summary {
    background: linear-gradient(135deg, #ffffff, #f8faff);
    border: 1px solid #c7d2fe;
    border-radius: 16px;
    padding: 17px 19px;
    margin-bottom: 14px;
}

.rubric-preview-title {
    color: #0f172a;
    font-size: 16px;
    font-weight: 780;
    line-height: 1.45;
    margin-bottom: 9px;
}

.rubric-preview-stats {
    display: flex;
    flex-wrap: wrap;
    gap: 8px;
}

.rubric-preview-stat {
    display: inline-flex;
    align-items: center;
    padding: 6px 10px;
    border-radius: 999px;
    background: #ffffff;
    border: 1px solid #e2e8f0;
    color: #475569;
    font-size: 11.5px;
    font-weight: 650;
}

.rubric-preview-heading {
    color: #475569;
    font-size: 11px;
    font-weight: 780;
    letter-spacing: 0.9px;
    text-transform: uppercase;
    margin: 15px 0 8px 0;
}

.rubric-preview-list {
    border: 1px solid #e2e8f0;
    border-radius: 15px;
    overflow: hidden;
    background: #ffffff;
}

.rubric-preview-item {
    display: grid;
    grid-template-columns: 36px minmax(0, 1fr) auto;
    gap: 12px;
    align-items: start;
    padding: 14px 15px;
    border-bottom: 1px solid #edf1f5;
}

.rubric-preview-item:last-child {
    border-bottom: 0;
}

.rubric-preview-index {
    width: 30px;
    height: 30px;
    display: flex;
    align-items: center;
    justify-content: center;
    border-radius: 9px;
    background: #eef2ff;
    color: #4f46e5;
    font-size: 11px;
    font-weight: 800;
}

.rubric-preview-name {
    color: #0f172a;
    font-size: 13.5px;
    font-weight: 750;
    line-height: 1.45;
    margin-bottom: 4px;
}

.rubric-preview-description {
    color: #64748b;
    font-size: 12px;
    line-height: 1.55;
}

.rubric-preview-score {
    min-width: 90px;
    text-align: right;
    color: #4338ca;
    font-size: 12px;
    font-weight: 750;
    white-space: nowrap;
    padding-top: 2px;
}

.rubric-preview-actions-note {
    color: #64748b;
    font-size: 12px;
    line-height: 1.55;
    margin: 13px 0 8px 0;
}

.st-key-rubric_generation_scroll,
.st-key-rubric_edit_scroll {
    height: 68vh !important;
    max-height: 68vh !important;
    overflow-y: auto !important;
    overflow-x: hidden !important;
    padding-right: 8px !important;
    scrollbar-width: thin;
    scrollbar-color: #cbd5e1 transparent;
}

[class*="st-key-rubric_action_"] button,
[class*="st-key-rubric_modal_"] button {
    min-height: 40px !important;
    height: 40px !important;
    background: #ffffff !important;
    color: #4f46e5 !important;
    border: 1px solid #c7d2fe !important;
    box-shadow: none !important;
    font-size: 12px !important;
}

[class*="st-key-rubric_action_"] button:hover,
[class*="st-key-rubric_modal_"] button:hover {
    background: #f5f7ff !important;
    color: #4338ca !important;
    border-color: #818cf8 !important;
}

.st-key-rubric_modal_use button,
.st-key-save_rubric_changes button {
    background: #4f46e5 !important;
    color: #ffffff !important;
    border-color: #4f46e5 !important;
}

.st-key-rubric_modal_use button:hover,
.st-key-save_rubric_changes button:hover {
    background: #4338ca !important;
    color: #ffffff !important;
    border-color: #4338ca !important;
}

@media (max-width: 700px) {
    .rubric-preview-item {
        grid-template-columns: 32px minmax(0, 1fr);
    }

    .rubric-preview-score {
        grid-column: 2;
        text-align: left;
        padding-top: 0;
    }
}

/* =====================================================
   EVALUATION MODAL
===================================================== */

div[data-testid="stDialog"] {
    padding: 12px !important;
    overflow: hidden !important;
}

div[data-testid="stDialog"] div[role="dialog"] {
    width: min(1120px, 92vw) !important;
    max-width: 92vw !important;
    max-height: 88vh !important;
    margin: auto !important;
    border-radius: 18px !important;
    overflow: hidden !important;
    box-shadow: 0 24px 70px rgba(15, 23, 42, 0.20) !important;
}

.st-key-evaluation_modal_scroll {
    height: 72vh !important;
    max-height: 72vh !important;
    overflow-y: auto !important;
    overflow-x: hidden !important;
    padding-right: 8px !important;
    scrollbar-width: thin;
    scrollbar-color: #cbd5e1 transparent;
}

.st-key-evaluation_modal_scroll::-webkit-scrollbar {
    width: 8px;
}

.st-key-evaluation_modal_scroll::-webkit-scrollbar-track {
    background: transparent;
}

.st-key-evaluation_modal_scroll::-webkit-scrollbar-thumb {
    background: #cbd5e1;
    border-radius: 999px;
}

.modal-header {
    margin-bottom: 14px;
}

.modal-title {
    color: #0f172a;
    font-size: 24px;
    font-weight: 780;
    letter-spacing: -0.45px;
    margin-bottom: 7px;
}

.modal-description {
    color: #64748b;
    font-size: 13.5px;
    line-height: 1.65;
    max-width: 760px;
}

.evaluation-loading-card {
    background: linear-gradient(135deg, #ffffff, #f7f8ff);
    border: 1px solid #c7d2fe;
    border-radius: 17px;
    padding: 25px;
    margin: 10px 0;
    box-shadow: 0 10px 30px rgba(79, 70, 229, 0.08);
}

.evaluation-loading-title {
    color: #0f172a;
    font-size: 18px;
    font-weight: 760;
    margin-bottom: 8px;
}

.evaluation-loading-text {
    color: #64748b;
    font-size: 13.5px;
    line-height: 1.65;
    margin-bottom: 18px;
}

.evaluation-loading-track {
    position: relative;
    width: 100%;
    height: 7px;
    background: #e8eaf5;
    border-radius: 999px;
    overflow: hidden;
}

.evaluation-loading-bar {
    position: absolute;
    top: 0;
    left: -40%;
    width: 38%;
    height: 100%;
    background: linear-gradient(90deg, #4f46e5, #6366f1, #818cf8);
    border-radius: 999px;
    animation: evaluation-loading 1.2s ease-in-out infinite;
}

@keyframes evaluation-loading {
    0% { left: -40%; }
    50% { left: 45%; }
    100% { left: 105%; }
}

.evaluation-success-card {
    background: #ffffff;
    border: 1px solid #d8e0ea;
    border-radius: 16px;
    padding: 20px 22px;
    margin: 10px 0 18px 0;
    box-shadow: 0 7px 20px rgba(15, 23, 42, 0.04);
}

.evaluation-success-title {
    color: #0f172a;
    font-size: 17px;
    font-weight: 760;
    margin-bottom: 6px;
}

.evaluation-success-text {
    color: #64748b;
    font-size: 13px;
    line-height: 1.65;
}

.evaluation-error-card {
    background: #fffafa;
    border: 1px solid #fecaca;
    border-radius: 16px;
    padding: 20px 22px;
    margin: 10px 0;
}

.evaluation-error-title {
    color: #991b1b;
    font-size: 15px;
    font-weight: 750;
    margin-bottom: 6px;
}

.evaluation-error-text {
    color: #7f1d1d;
    font-size: 13px;
    line-height: 1.6;
}

.evaluation-download-card {
    background: linear-gradient(135deg, #f8faff, #ffffff);
    border: 1px solid #c7d2fe;
    border-radius: 15px;
    padding: 16px 18px;
    margin-top: 18px;
}

.evaluation-download-title {
    color: #0f172a;
    font-size: 14px;
    font-weight: 760;
    margin-bottom: 4px;
}

.evaluation-download-text {
    color: #64748b;
    font-size: 12.5px;
    line-height: 1.55;
}

.st-key-download_evaluation_report button {
    min-height: 42px !important;
    height: 42px !important;
    margin-top: 8px !important;
    background: #4f46e5 !important;
    color: #ffffff !important;
    border: 1px solid #4f46e5 !important;
    border-radius: 11px !important;
    font-size: 12.5px !important;
    font-weight: 700 !important;
    box-shadow: 0 7px 16px rgba(79, 70, 229, 0.14) !important;
}

.st-key-download_evaluation_report button:hover {
    background: #4338ca !important;
    color: #ffffff !important;
    border-color: #4338ca !important;
}

.st-key-evaluation_report_format {
    margin-top: 10px !important;
}

.st-key-evaluation_report_format label {
    color: #475569 !important;
    font-size: 12px !important;
    font-weight: 700 !important;
}

.st-key-evaluation_report_format [data-baseweb="select"] > div {
    min-height: 42px !important;
    border-radius: 11px !important;
    border-color: #cbd5e1 !important;
    background: #ffffff !important;
}

@media (max-height: 800px) {
    .st-key-evaluation_modal_scroll {
        height: 67vh !important;
        max-height: 67vh !important;
    }
}

@media (max-width: 900px) {
    .block-container {
        padding-top: 30px !important;
        padding-left: 20px !important;
        padding-right: 20px !important;
    }

    .page-title {
        font-size: 31px;
    }
}

</style>
"""
)



# =========================================================
# RUBRIC GENERATION MODAL
# =========================================================

@st.dialog("AI Rubric Generation", width="large")
def rubric_generation_modal():
    with st.container(key="rubric_generation_scroll"):
        phase = st.session_state.get(
            "rubric_generation_phase",
            "idle",
        )

        if phase == "ready":
            modal_title = "Review AI-Generated Rubric"
            modal_description = (
                "Review the generated criteria, marks, and weightage before "
                "using this rubric for assignment evaluation."
            )
        else:
            modal_title = "Generate Rubric from Assignment"
            modal_description = (
                "EduAI is analysing the uploaded assignment and creating "
                "a structured grading rubric."
            )

        st.html(
            f"""
            <div class="modal-header">
                <div class="modal-title">{safe_html(modal_title)}</div>
                <div class="modal-description">
                    {safe_html(modal_description)}
                </div>
            </div>
            """
        )

        snapshot = st.session_state.get(
            "rubric_assignment_snapshot"
        )

        if not snapshot:
            st.error("Assignment file is not available.")
            return

        if st.session_state.get(
            "rubric_generation_phase"
        ) == "generating":
            loading = st.empty()

            loading.html(
                """
                <div class="evaluation-loading-card">
                    <div class="evaluation-loading-title">
                        Generating rubric
                    </div>
                    <div class="evaluation-loading-text">
                        Reading assignment requirements, identifying measurable
                        criteria, allocating marks, and validating the rubric
                        structure.
                    </div>
                    <div class="evaluation-loading-track">
                        <div class="evaluation-loading-bar"></div>
                    </div>
                </div>
                """
            )

            try:
                result = generate_rubric_from_assignment(
                    snapshot
                )

                rubric = normalize_rubric(
                    result
                )

                valid, message = validate_rubric(
                    rubric
                )

                if not valid:
                    raise ValueError(
                        message
                    )

                st.session_state.generated_rubric = rubric
                st.session_state.rubric_source = "ai"
                st.session_state.rubric_generation_error = None
                st.session_state.rubric_generation_phase = "ready"

            except Exception as error:
                st.session_state.rubric_generation_error = str(
                    error
                )
                st.session_state.rubric_generation_phase = "error"

            loading.empty()

        phase = st.session_state.get(
            "rubric_generation_phase"
        )

        if phase == "error":
            st.error(
                st.session_state.get(
                    "rubric_generation_error"
                )
                or
                "Rubric generation failed."
            )

            if st.button(
                "Try Again",
                key="rubric_generation_retry",
                use_container_width=True,
            ):
                st.session_state.rubric_generation_phase = "generating"
                st.session_state.rubric_generation_error = None
                st.rerun(
                    scope="fragment"
                )

            return

        rubric = st.session_state.get(
            "generated_rubric"
        )

        if phase == "ready" and rubric:
            criteria = rubric.get(
                "criteria",
                [],
            )

            total_marks = safe_number(
                rubric.get(
                    "total_marks",
                    0,
                )
            )

            criteria_total = sum(
                safe_number(
                    criterion.get(
                        "max_marks",
                        0,
                    )
                )
                for criterion in criteria
            )

            st.html(
                """
                <div class="rubric-preview-success">
                    Rubric generated successfully. Review it before continuing.
                </div>
                """
            )

            st.html(
                f"""
                <div class="rubric-preview-summary">
                    <div class="rubric-source-badge">AI Generated</div>
                    <div class="rubric-preview-title">
                        {safe_html(rubric.get('title', 'AI Generated Rubric'))}
                    </div>
                    <div class="rubric-preview-stats">
                        <div class="rubric-preview-stat">
                            {len(criteria)} Criteria
                        </div>
                        <div class="rubric-preview-stat">
                            Total Marks: {format_score(total_marks)}
                        </div>
                        <div class="rubric-preview-stat">
                            Criteria Total: {format_score(criteria_total)}
                        </div>
                    </div>
                </div>
                """
            )

            preview_items = []

            for index, criterion in enumerate(
                criteria,
                start=1,
            ):
                criterion_name = safe_html(
                    criterion.get(
                        "name",
                        f"Criterion {index}",
                    )
                )

                description = safe_html(
                    criterion.get(
                        "description",
                        "",
                    )
                )

                marks = safe_number(
                    criterion.get(
                        "max_marks",
                        0,
                    )
                )

                weight = safe_number(
                    criterion.get(
                        "weight",
                        0,
                    )
                )

                preview_items.append(
                    f"""
                    <div class="rubric-preview-item">
                        <div class="rubric-preview-index">
                            {index:02d}
                        </div>
                        <div>
                            <div class="rubric-preview-name">
                                {criterion_name}
                            </div>
                            <div class="rubric-preview-description">
                                {description or 'No description provided.'}
                            </div>
                        </div>
                        <div class="rubric-preview-score">
                            {format_score(marks)} Marks<br>
                            {weight:.2f}%
                        </div>
                    </div>
                    """
                )

            st.html(
                f"""
                <div class="rubric-preview-heading">
                    Rubric Preview
                </div>
                <div class="rubric-preview-list">
                    {''.join(preview_items)}
                </div>
                <div class="rubric-preview-actions-note">
                    Confirm this rubric to continue, edit the generated criteria,
                    or regenerate a fresh rubric from the same assignment.
                </div>
                """
            )

            edit_col, regenerate_col, use_col = st.columns(
                [1, 1, 1.25],
                gap="small",
            )

            with edit_col:
                if st.button(
                    "Edit Rubric",
                    key="rubric_modal_edit",
                    use_container_width=True,
                ):
                    st.session_state.open_rubric_edit_after_generation = True
                    st.rerun()

            with regenerate_col:
                if st.button(
                    "Regenerate",
                    key="rubric_modal_regenerate",
                    use_container_width=True,
                ):
                    st.session_state.rubric_generation_phase = "generating"
                    st.session_state.rubric_generation_error = None
                    st.session_state.rubric_source = "ai"
                    st.session_state.uploaded_rubric_snapshot = None
                    st.session_state.show_rubric_upload = False
                    st.rerun(
                        scope="fragment"
                    )

            with use_col:
                if st.button(
                    "Use This Rubric",
                    key="rubric_modal_use",
                    use_container_width=True,
                ):
                    st.session_state.rubric_source = "ai"
                    st.session_state.uploaded_rubric_snapshot = None
                    st.session_state.show_rubric_upload = False
                    st.rerun()


# =========================================================
# RUBRIC EDIT MODAL
# =========================================================

@st.dialog("Edit AI Rubric", width="large")
def rubric_edit_modal():
    rubric=st.session_state.get("generated_rubric")
    if not rubric:
        st.error("No AI-generated rubric is available to edit.")
        return

    with st.container(key="rubric_edit_scroll"):
        st.html("""
        <div class="modal-header">
            <div class="modal-title">Edit Rubric</div>
            <div class="modal-description">
                Review the generated criteria and marks. Save once after completing all changes.
            </div>
        </div>
        """)
        title=st.text_input("Rubric Title",value=str(rubric.get("title","AI Generated Rubric")),key="edit_rubric_title")
        total=st.number_input("Total Marks",min_value=1.0,value=float(rubric.get("total_marks",1)),step=1.0,key="edit_rubric_total_marks")
        edited=[]
        for i,c in enumerate(rubric.get("criteria",[]),start=1):
            with st.container(border=True):
                st.html(f'<div class="section-number">Criterion {i}</div>')
                name=st.text_input("Criterion Name",value=str(c.get("name","")),key=f"rubric_edit_name_{i}")
                desc=st.text_area("Description",value=str(c.get("description","")),height=100,key=f"rubric_edit_description_{i}")
                marks=st.number_input("Marks",min_value=0.0,value=float(c.get("max_marks",0)),step=1.0,key=f"rubric_edit_marks_{i}")
                edited.append({"name":name.strip(),"description":desc.strip(),"max_marks":marks,"levels":c.get("levels",[])})
        marks_total=sum(safe_number(x.get("max_marks",0)) for x in edited)
        st.html(f'<div class="rubric-total-bar">Criteria Total: <strong>{format_score(marks_total)}</strong> &nbsp;/&nbsp; Rubric Total: <strong>{format_score(total)}</strong></div>')
        if st.button("Save Changes",key="save_rubric_changes",use_container_width=True):
            if not title.strip():
                st.warning("Rubric title cannot be empty.")
            elif abs(marks_total-total)>0.05:
                st.error("Criterion marks must equal the rubric total marks.")
            elif any(not x["name"] for x in edited):
                st.warning("Every rubric criterion must have a name.")
            else:
                for x in edited:
                    x["weight"]=(safe_number(x["max_marks"])/total*100) if total else 0
                st.session_state.generated_rubric={"title":title.strip(),"total_marks":total,"criteria":edited}
                st.session_state.rubric_source="ai"
                st.success("Rubric changes saved successfully.")


# =========================================================
# EVALUATION LOGS MODAL
# =========================================================

@st.dialog(
    "Evaluation Logs",
    width="large",
)
def evaluation_logs_modal():

    with st.container(
        key="ai_eval_logs_scroll",
    ):

        st.html(
        """
        <div class="eval-logs-header">
            <div class="eval-logs-title">
                AI Assignment Evaluation Logs
            </div>
            <div class="eval-logs-description">
                Review assignment evaluations completed during
                the current application session.
            </div>
        </div>
        """
        )


        logs = st.session_state.get(
            "ai_evaluation_logs",
            [],
        )


        st.html(
            f"""
            <div class="eval-log-summary">
                <strong>{len(logs)}</strong>
                evaluation record(s) available.
            </div>
            """
        )


        if not logs:

            st.info(
                "No assignment evaluations have been completed in this session yet."
            )

            return


        search = st.text_input(
            "Search logs",
            placeholder="Search by assignment or submission file",
            key="ai_evaluation_log_search",
        )


        filtered_logs = logs


        if search:

            search_value = search.strip().lower()


            filtered_logs = [
                log
                for log in logs
                if (
                    search_value
                    in str(
                        log.get(
                            "assignment_name",
                            "",
                        )
                    ).lower()
                    or
                    search_value
                    in str(
                        log.get(
                            "submission_file",
                            "",
                        )
                    ).lower()
                )
            ]


        if not filtered_logs:

            st.info(
                "No evaluation logs match the current search."
            )

            return


        for index, log in enumerate(
            reversed(
                filtered_logs
            ),
            start=1,
        ):

            assignment_name = safe_html(
                log.get(
                    "assignment_name",
                    "Assignment Evaluation",
                )
            )


            evaluated_at = safe_html(
                log.get(
                    "evaluated_at",
                    "-",
                )
            )


            assignment_file = safe_html(
                log.get(
                    "assignment_file",
                    "-",
                )
            )


            rubric_file = safe_html(
                log.get(
                    "rubric_file",
                    "-",
                )
            )


            submission_file = safe_html(
                log.get(
                    "submission_file",
                    "-",
                )
            )


            total_marks = safe_number(
                log.get(
                    "total_marks",
                    0,
                )
            )


            obtained_marks = safe_number(
                log.get(
                    "obtained_marks",
                    0,
                )
            )


            percentage = safe_number(
                log.get(
                    "percentage",
                    0,
                )
            )


            st.html(
                f"""
                <div class="eval-log-card">
                    <div class="eval-log-title">
                        {assignment_name}
                    </div>

                    <div class="eval-log-meta">
                        Evaluated: {evaluated_at}<br>
                        Assignment: {assignment_file}<br>
                        Rubric: {rubric_file}<br>
                        Submission: {submission_file}
                    </div>

                    <div class="eval-log-score">
                        Score:
                        {format_score(obtained_marks)}
                        /
                        {format_score(total_marks)}
                        &nbsp;&nbsp; | &nbsp;&nbsp;
                        {percentage:.2f}%
                    </div>
                </div>
                """
            )


# =========================================================
# PAGE HEADER
# =========================================================

health = check_assignment_evaluation_health()
is_healthy = health.get("status") == "healthy"


back_col, header_col, logs_col = st.columns(
    [
        0.55,
        8.05,
        1.40,
    ],
    gap="small",
    vertical_alignment="top",
)


with back_col:

    if st.button(
        "‹",
        key="ai_eval_back_dashboard",
        help="Back to dashboard",
    ):
        st.switch_page(
            "pages/instructor_dashboard.py"
        )


with header_col:

    st.html(
        f"""
    <div class="page-header">
        <div class="page-label">EduAI Assessment</div>
        <div class="page-title">AI Assignment Evaluation</div>
        <div class="page-description">
            Upload an assignment and EduAI will generate a grading rubric
            automatically. Review or replace the rubric, then upload the student's
            submission and run the final AI evaluation.
        </div>
        <div class="status-row">
            <span class="status-dot {' ' if is_healthy else 'offline'}"></span>
            {'Evaluation backend connected' if is_healthy else 'Evaluation backend unavailable'}
        </div>
    </div>
    """
    )


with logs_col:

    st.markdown(
        "<div style='height:28px;'></div>",
        unsafe_allow_html=True,
    )


    if st.button(
        "Generated Records",
        key="ai_eval_view_logs",
        use_container_width=True,
    ):

        evaluation_logs_modal()


accepted_types = [
    "pdf",
    "docx",
    "txt",
    "py",
    "java",
    "cpp",
    "c",
    "html",
    "css",
    "js",
    "json",
    "sql",
    "md",
]


# =========================================================
# STEP 01 — ASSIGNMENT
# =========================================================

assignment_file = None
selected_question_bank_set = None

with st.container(
    border=True
):

    st.html(
        """
        <div class="section-header">
            <div class="section-number">
                Step 01
            </div>

            <div class="section-title">
                Assignment
            </div>

            <div class="section-description">
                Upload an assignment file or select a previously saved
                assessment from the Question Bank. EduAI will use the
                selected assignment to generate the grading rubric.
            </div>
        </div>
        """
    )

    assignment_source = st.radio(
        "Assignment Source",
        [
            "Upload Assignment",
            "Select from Question Bank",
        ],
        horizontal=True,
        key="ai_assignment_source",
        label_visibility="collapsed",
    )

    st.markdown(
        ""
    )

    # -----------------------------------------------------
    # UPLOAD ASSIGNMENT
    # -----------------------------------------------------

    if (
        assignment_source
        == "Upload Assignment"
    ):

        assignment_file = st.file_uploader(
            "Assignment",
            type=accepted_types,
            key="ai_assignment_file",
            label_visibility="collapsed",
        )

        st.html(
            (
                '<div class="file-pill '
                + (
                    'ready'
                    if assignment_file
                    else
                    ''
                )
                + '">'
                + (
                    f"Selected: {safe_html(assignment_file.name)}"
                    if assignment_file
                    else
                    "Assignment not selected"
                )
                + '</div>'
            )
        )

    # -----------------------------------------------------
    # QUESTION BANK
    # -----------------------------------------------------

    else:

        try:

            question_bank_sets = (
                get_question_sets()
                or
                []
            )

        except Exception as question_bank_error:

            question_bank_sets = []

            st.error(
                "Unable to load the Question Bank. "
                "Please check the backend connection."
            )

            print(
                "Question Bank load error:",
                question_bank_error,
            )

        if not question_bank_sets:

            st.info(
                "No saved Question Bank assignments are available yet. "
                "Generate and save a question set first."
            )

        else:

            normalized_bank_sets = []

            for item in question_bank_sets:

                if not isinstance(
                    item,
                    dict,
                ):
                    continue

                normalized_item = item.copy()

                item_id = (
                    normalized_item.get(
                        "id"
                    )
                    or
                    normalized_item.get(
                        "set_id"
                    )
                )

                if item_id is None:
                    continue

                normalized_item[
                    "id"
                ] = item_id

                normalized_bank_sets.append(
                    normalized_item
                )

            if not normalized_bank_sets:

                st.info(
                    "Question Bank records were found, but none contain "
                    "a valid question-set ID."
                )

            else:

                bank_options = {
                    question_bank_option_label(
                        item
                    ):
                    item
                    for item in normalized_bank_sets
                }

                selected_bank_label = st.selectbox(
                    "Select Assignment from Question Bank",
                    options=list(
                        bank_options.keys()
                    ),
                    index=None,
                    placeholder=(
                        "Choose a saved assignment"
                    ),
                    key="ai_question_bank_assignment",
                )

                if selected_bank_label:

                    selected_summary = (
                        bank_options[
                            selected_bank_label
                        ]
                    )

                    selected_id = (
                        selected_summary.get(
                            "id"
                        )
                    )

                    try:

                        (
                            assignment_file,
                            selected_question_bank_set,
                            bank_file_format,
                        ) = (
                            build_question_bank_assignment_file(
                                question_set_id=
                                    selected_id,
                                question_set_summary=
                                    selected_summary,
                            )
                        )

                        selected_title = (
                            selected_question_bank_set.get(
                                "title",
                                "Question Bank Assignment",
                            )
                        )

                        selected_topic = (
                            selected_question_bank_set.get(
                                "topic",
                                "-",
                            )
                        )

                        selected_difficulty = (
                            selected_question_bank_set.get(
                                "difficulty",
                                "-",
                            )
                        )

                        selected_total_marks = (
                            selected_question_bank_set.get(
                                "total_marks",
                                "-",
                            )
                        )

                        st.html(
                            f"""
                            <div class="question-bank-card">
                                <div class="question-bank-title">
                                    {safe_html(selected_title)}
                                </div>

                                <div class="question-bank-meta">
                                    Topic:
                                    <strong>
                                        {safe_html(selected_topic)}
                                    </strong>

                                    &nbsp;&nbsp; | &nbsp;&nbsp;

                                    Difficulty:
                                    <strong>
                                        {safe_html(selected_difficulty)}
                                    </strong>

                                    &nbsp;&nbsp; | &nbsp;&nbsp;

                                    Total Marks:
                                    <strong>
                                        {safe_html(selected_total_marks)}
                                    </strong>

                                    &nbsp;&nbsp; | &nbsp;&nbsp;

                                    Source:
                                    <strong>
                                        Question Bank
                                    </strong>
                                </div>
                            </div>
                            """
                        )

                        st.html(
                            (
                                '<div class="file-pill ready">'
                                'Selected from Question Bank: '
                                + safe_html(
                                    assignment_file.name
                                )
                                + '</div>'
                            )
                        )

                        if (
                            bank_file_format
                            == "txt"
                        ):

                            st.caption(
                                "The saved Question Bank PDF export was "
                                "unavailable, so EduAI reconstructed the "
                                "assignment from the saved questions."
                            )

                    except Exception as selected_bank_error:

                        assignment_file = None

                        st.error(
                            "Unable to prepare the selected Question Bank "
                            "assignment. Please try another saved assessment."
                        )

                        print(
                            "Question Bank assignment error:",
                            selected_bank_error,
                        )


auto_open_rubric_modal=False
if assignment_file is not None:
    current_snapshot=create_file_snapshot(assignment_file)
    current_hash=calculate_file_hash(current_snapshot)
    if current_hash != st.session_state.get("rubric_assignment_hash"):
        reset_rubric_flow()
        st.session_state.rubric_assignment_hash=current_hash
        st.session_state.rubric_assignment_snapshot=current_snapshot
        st.session_state.rubric_generation_phase="generating"
        auto_open_rubric_modal=True
else:
    if st.session_state.get("rubric_assignment_hash") is not None:
        st.session_state.rubric_assignment_hash=None
        st.session_state.rubric_assignment_snapshot=None
        reset_rubric_flow()

if auto_open_rubric_modal:
    rubric_generation_modal()


if st.session_state.get(
    "open_rubric_edit_after_generation"
):
    st.session_state.open_rubric_edit_after_generation = False
    rubric_edit_modal()


# =========================================================
# STEP 02 — RUBRIC
# =========================================================

active_rubric_snapshot=None
if assignment_file is not None:
    with st.container(border=True):
        st.html("""
        <div class="section-header">
            <div class="section-number">Step 02</div>
            <div class="section-title">Rubric</div>
            <div class="section-description">
                Review the AI-generated rubric, edit it, regenerate it, or replace it with your own rubric file.
            </div>
        </div>
        """)
        rubric=st.session_state.get("generated_rubric")
        phase=st.session_state.get("rubric_generation_phase")
        if phase == "generating":
            st.info("AI rubric generation is in progress.")
        elif phase == "error":
            st.error(st.session_state.get("rubric_generation_error") or "Rubric generation failed.")
            if st.button("Generate Again",key="rubric_action_retry"):
                st.session_state.rubric_generation_phase="generating"
                rubric_generation_modal()
        elif rubric:
            source = (
                st.session_state.get("rubric_source")
                or "ai"
            )

            if source == "uploaded":
                uploaded_snapshot = (
                    st.session_state.get(
                        "uploaded_rubric_snapshot"
                    )
                )

                uploaded_name = (
                    uploaded_snapshot.get(
                        "name",
                        "Instructor Rubric",
                    )
                    if uploaded_snapshot
                    else
                    "Instructor Rubric"
                )

                st.html(
                    f"""
                    <div class="rubric-ready-card">
                        <div class="rubric-source-badge">
                            Instructor Uploaded
                        </div>
                        <div class="rubric-ready-title">
                            {safe_html(uploaded_name)}
                        </div>
                        <div class="rubric-ready-meta">
                            This uploaded rubric will be used for
                            the final assignment evaluation.
                        </div>
                    </div>
                    """
                )

                c1, c2, c3 = st.columns(
                    3,
                    gap="small",
                )

                with c1:
                    if st.button(
                        "Use AI Rubric",
                        key="rubric_action_use_ai",
                        use_container_width=True,
                    ):
                        st.session_state.rubric_source = "ai"
                        st.session_state.uploaded_rubric_snapshot = None
                        st.session_state.show_rubric_upload = False
                        st.rerun()

                with c2:
                    if st.button(
                        "Upload Different Rubric",
                        key="rubric_action_upload",
                        use_container_width=True,
                    ):
                        st.session_state.show_rubric_upload = True

                with c3:
                    if st.button(
                        "Regenerate AI Rubric",
                        key="rubric_action_regenerate",
                        use_container_width=True,
                    ):
                        st.session_state.rubric_generation_phase = "generating"
                        st.session_state.rubric_generation_error = None
                        st.session_state.rubric_source = "ai"
                        st.session_state.uploaded_rubric_snapshot = None
                        st.session_state.show_rubric_upload = False
                        rubric_generation_modal()

            else:
                st.html(
                    f"""
                    <div class="rubric-ready-card">
                        <div class="rubric-source-badge">
                            AI Generated
                        </div>
                        <div class="rubric-ready-title">
                            {safe_html(
                                rubric.get(
                                    'title',
                                    'AI Generated Rubric',
                                )
                            )}
                        </div>
                        <div class="rubric-ready-meta">
                            {len(rubric.get('criteria', []))}
                            criteria
                            &nbsp; | &nbsp;
                            Total Marks:
                            {format_score(
                                rubric.get(
                                    'total_marks',
                                    0,
                                )
                            )}
                        </div>
                    </div>
                    """
                )

                c1, c2, c3 = st.columns(
                    3,
                    gap="small",
                )

                with c1:
                    if st.button(
                        "Edit Rubric",
                        key="rubric_action_edit",
                        use_container_width=True,
                    ):
                        rubric_edit_modal()

                with c2:
                    if st.button(
                        "Upload Rubric",
                        key="rubric_action_upload",
                        use_container_width=True,
                    ):
                        st.session_state.show_rubric_upload = (
                            not st.session_state.get(
                                "show_rubric_upload"
                            )
                        )

                with c3:
                    if st.button(
                        "Regenerate",
                        key="rubric_action_regenerate",
                        use_container_width=True,
                    ):
                        st.session_state.rubric_generation_phase = "generating"
                        st.session_state.rubric_generation_error = None
                        st.session_state.rubric_source = "ai"
                        st.session_state.uploaded_rubric_snapshot = None
                        st.session_state.show_rubric_upload = False
                        rubric_generation_modal()

            if st.session_state.get(
                "show_rubric_upload"
            ):
                st.markdown("")

                custom = st.file_uploader(
                    "Upload Your Rubric",
                    type=[
                        "pdf",
                        "docx",
                        "txt",
                        "json",
                    ],
                    key="ai_custom_rubric_file",
                )

                if custom is not None:
                    st.session_state.uploaded_rubric_snapshot = (
                        create_file_snapshot(
                            custom
                        )
                    )

                    st.session_state.rubric_source = "uploaded"

                    st.success(
                        f"Using uploaded rubric: {custom.name}"
                    )

        active_rubric_snapshot = get_active_rubric_snapshot()


# =========================================================
# STEP 03 — STUDENT SUBMISSION
# =========================================================

submission_files = []
submission_selection_valid = False
evaluate_clicked = False

if assignment_file is not None and active_rubric_snapshot:

    with st.container(
        border=True
    ):

        st.html(
            """
            <div class="section-header">
                <div class="section-number">Step 03</div>
                <div class="section-title">Student Submission</div>
                <div class="section-description">
                    Upload one student PDF/DOCX file, multiple student files,
                    or one ZIP containing student submissions.
                    Every submission is evaluated separately against the same
                    active assignment and rubric.
                </div>
            </div>
            """
        )

        submission_files = st.file_uploader(
            "Student Submission",
            type=[
                "pdf",
                "docx",
                "zip",
            ],
            accept_multiple_files=True,
            key="ai_submission_file",
            label_visibility="collapsed",
        )

        submission_validation = (
            validate_submission_selection(
                submission_files
            )
        )

        submission_selection_valid = bool(
            submission_validation.get(
                "valid"
            )
        )

        submission_label = (
            submission_selection_label(
                submission_files
            )
        )

        st.html(
            (
                '<div class="file-pill '
                + (
                    'ready'
                    if submission_selection_valid
                    else ''
                )
                + '">'
                + safe_html(
                    submission_label
                )
                + '</div>'
            )
        )

        if (
            submission_files
            and
            not submission_selection_valid
        ):

            st.warning(
                submission_validation.get(
                    "message"
                )
                or
                "Please check the selected submission files."
            )

        elif (
            submission_selection_valid
            and
            submission_validation.get(
                "mode"
            ) == "zip"
        ):

            st.caption(
                "Each PDF inside the ZIP will be treated as a separate "
                "student submission and will receive its own evaluation report."
            )

        elif (
            submission_selection_valid
            and
            len(
                submission_files
            ) > 1
        ):

            st.caption(
                f"{len(submission_files)} student PDFs will be evaluated "
                "separately using the same active rubric."
            )

        st.info(
            "Batch rule: all uploaded student files must belong to the "
            "same assignment shown in Step 01. The active rubric will be "
            "applied independently to every student submission."
        )

        st.markdown(
            ""
        )

        evaluate_clicked = st.button(
            "Evaluate Assignment  →",
            use_container_width=True,
            key="run_ai_assignment_evaluation",
            disabled=(
                not submission_selection_valid
            ),
        )


# =========================================================
# EVALUATION MODAL
# =========================================================

@st.dialog(
    "Assignment Evaluation",
    width="large",
)
def evaluation_modal():

    with st.container(
        key="evaluation_modal_scroll",
    ):

        snapshot = st.session_state.get(
            "ai_evaluation_snapshot"
        )

        if not snapshot:

            st.error(
                "Evaluation files are unavailable. "
                "Close this window and try again."
            )

            return

        assignment_snapshot = snapshot.get(
            "assignment_file"
        )

        rubric_snapshot = snapshot.get(
            "rubric_file"
        )

        submission_snapshots = snapshot.get(
            "submission_files",
            [],
        )

        if not assignment_snapshot:
            st.warning(
                "Please upload the assignment file."
            )
            return

        if not rubric_snapshot:
            st.warning(
                "A generated or uploaded rubric is required."
            )
            return

        if not submission_snapshots:
            st.warning(
                "Please upload at least one student submission."
            )
            return

        batch_count = len(
            submission_snapshots
        )

        st.html(
            f"""
            <div class="modal-header">
                <div class="modal-title">
                    AI Assignment Evaluation
                </div>
                <div class="modal-description">
                    {batch_count} student submission(s) will be evaluated
                    independently against the same active rubric.
                    Each student receives an individual result and report.
                </div>
            </div>
            """
        )

        # -------------------------------------------------
        # RUN BATCH EVALUATION
        # -------------------------------------------------

        if (
            st.session_state.get(
                "ai_evaluation_phase"
            )
            == "evaluating"
        ):

            loading_area = st.empty()
            batch_results = []

            for index, submission_snapshot in enumerate(
                submission_snapshots,
                start=1,
            ):

                submission_name = submission_snapshot.get(
                    "name",
                    f"Student_Submission_{index}.pdf",
                )

                loading_area.html(
                    f"""
                    <div class="evaluation-loading-card">
                        <div class="evaluation-loading-title">
                            Evaluating student {index} of {batch_count}
                        </div>
                        <div class="evaluation-loading-text">
                            {safe_html(submission_name)} is being evaluated
                            against the active rubric.
                        </div>
                        <div class="evaluation-loading-track">
                            <div class="evaluation-loading-bar"></div>
                        </div>
                    </div>
                    """
                )

                try:

                    api_result = evaluate_assignment(
                        assignment_file=rebuild_uploaded_file(
                            assignment_snapshot
                        ),
                        rubric_file=rebuild_uploaded_file(
                            rubric_snapshot
                        ),
                        submission_file=rebuild_uploaded_file(
                            submission_snapshot
                        ),
                    )

                    evaluation = get_evaluation_payload(
                        api_result
                    )

                    if (
                        not evaluation
                        or evaluation.get(
                            "error"
                        )
                    ):

                        error_message = (
                            evaluation.get(
                                "error",
                                "Evaluation failed.",
                            )
                            if evaluation
                            else
                            "No valid evaluation result was returned."
                        )

                        details = (
                            evaluation.get(
                                "details"
                            )
                            if evaluation
                            else
                            None
                        )

                        if details:
                            error_message = (
                                f"{error_message} {details}"
                            )

                        batch_results.append(
                            {
                                "submission_snapshot":
                                    submission_snapshot,
                                "api_result":
                                    api_result,
                                "evaluation":
                                    None,
                                "error":
                                    error_message,
                            }
                        )

                        continue

                    batch_results.append(
                        {
                            "submission_snapshot":
                                submission_snapshot,
                            "api_result":
                                api_result,
                            "evaluation":
                                evaluation,
                            "error":
                                None,
                        }
                    )

                    # One log per student.
                    total_marks_for_log = safe_number(
                        evaluation.get(
                            "max_score",
                            evaluation.get(
                                "total_marks",
                                0,
                            ),
                        )
                    )

                    obtained_marks_for_log = safe_number(
                        evaluation.get(
                            "total_score",
                            evaluation.get(
                                "obtained_marks",
                                0,
                            ),
                        )
                    )

                    percentage_for_log = (
                        (
                            obtained_marks_for_log
                            /
                            total_marks_for_log
                        )
                        * 100
                        if total_marks_for_log > 0
                        else 0.0
                    )

                    st.session_state.ai_evaluation_logs.append(
                        {
                            "assignment_name":
                                evaluation.get(
                                    "assignment_name"
                                )
                                or
                                assignment_snapshot.get(
                                    "name",
                                    "Assignment Evaluation",
                                ),

                            "assignment_file":
                                assignment_snapshot.get(
                                    "name",
                                    "-",
                                ),

                            "rubric_file":
                                rubric_snapshot.get(
                                    "name",
                                    "-",
                                ),

                            "submission_file":
                                submission_snapshot.get(
                                    "name",
                                    "-",
                                ),

                            "total_marks":
                                total_marks_for_log,

                            "obtained_marks":
                                obtained_marks_for_log,

                            "percentage":
                                percentage_for_log,

                            "evaluated_at":
                                datetime.now().strftime(
                                    "%Y-%m-%d %H:%M:%S"
                                ),
                        }
                    )

                except Exception as error:

                    batch_results.append(
                        {
                            "submission_snapshot":
                                submission_snapshot,
                            "api_result":
                                None,
                            "evaluation":
                                None,
                            "error":
                                str(
                                    error
                                ),
                        }
                    )

            loading_area.empty()

            st.session_state.ai_evaluation_batch_results = (
                batch_results
            )

            successful_results = [
                item
                for item in batch_results
                if item.get(
                    "evaluation"
                )
            ]

            if successful_results:

                st.session_state.ai_evaluation_result = (
                    successful_results[0].get(
                        "api_result"
                    )
                )

                st.session_state.ai_evaluation_error = None

                st.session_state.ai_evaluation_phase = (
                    "result"
                )

            else:

                st.session_state.ai_evaluation_error = (
                    "None of the student submissions could be evaluated."
                )

                st.session_state.ai_evaluation_phase = (
                    "error"
                )

        # -------------------------------------------------
        # COMPLETE FAILURE
        # -------------------------------------------------

        if (
            st.session_state.get(
                "ai_evaluation_phase"
            )
            == "error"
        ):

            st.html(
                f"""
                <div class="evaluation-error-card">
                    <div class="evaluation-error-title">
                        Evaluation could not be completed
                    </div>
                    <div class="evaluation-error-text">
                        {safe_html(
                            st.session_state.get(
                                "ai_evaluation_error"
                            )
                            or
                            "Evaluation failed."
                        )}
                    </div>
                </div>
                """
            )

            return

        if (
            st.session_state.get(
                "ai_evaluation_phase"
            )
            != "result"
        ):
            return

        batch_results = st.session_state.get(
            "ai_evaluation_batch_results",
            [],
        )

        successful_results = [
            item
            for item in batch_results
            if item.get(
                "evaluation"
            )
        ]

        failed_results = [
            item
            for item in batch_results
            if not item.get(
                "evaluation"
            )
        ]

        # -------------------------------------------------
        # BATCH SUMMARY
        # -------------------------------------------------

        st.html(
            f"""
            <div class="evaluation-success-card">
                <div class="evaluation-success-title">
                    Batch evaluation completed
                </div>
                <div class="evaluation-success-text">
                    {len(successful_results)} of {len(batch_results)}
                    student submission(s) were evaluated successfully
                    using the active rubric.
                </div>
            </div>
            """
        )

        summary1, summary2, summary3 = st.columns(
            3,
            gap="large",
        )

        summary1.metric(
            "Total Students",
            len(
                batch_results
            ),
        )

        summary2.metric(
            "Evaluated",
            len(
                successful_results
            ),
        )

        summary3.metric(
            "Failed",
            len(
                failed_results
            ),
        )

        # -------------------------------------------------
        # ONE RESULT + REPORT PER STUDENT
        # -------------------------------------------------

        st.markdown(
            "### Student Results"
        )

        for index, item in enumerate(
            batch_results,
            start=1,
        ):

            submission_snapshot = item.get(
                "submission_snapshot",
                {},
            )

            submission_name = submission_snapshot.get(
                "name",
                f"Student {index}",
            )

            evaluation = item.get(
                "evaluation"
            )

            error_message = item.get(
                "error"
            )

            if not evaluation:

                with st.expander(
                    f"{submission_name} — Evaluation Failed",
                    expanded=False,
                ):

                    st.error(
                        error_message
                        or
                        "This student submission could not be evaluated."
                    )

                continue

            total_marks = safe_number(
                evaluation.get(
                    "max_score",
                    evaluation.get(
                        "total_marks",
                        0,
                    ),
                )
            )

            obtained_marks = safe_number(
                evaluation.get(
                    "total_score",
                    evaluation.get(
                        "obtained_marks",
                        0,
                    ),
                )
            )

            percentage = (
                (
                    obtained_marks
                    /
                    total_marks
                )
                * 100
                if total_marks > 0
                else 0.0
            )

            with st.expander(
                (
                    f"{submission_name} — "
                    f"{format_score(obtained_marks)} / "
                    f"{format_score(total_marks)} "
                    f"({percentage:.2f}%)"
                ),
                expanded=(
                    len(
                        batch_results
                    )
                    == 1
                ),
            ):

                metric1, metric2, metric3 = st.columns(
                    3,
                    gap="large",
                )

                metric1.metric(
                    "Total Marks",
                    format_score(
                        total_marks
                    ),
                )

                metric2.metric(
                    "Marks Obtained",
                    format_score(
                        obtained_marks
                    ),
                )

                metric3.metric(
                    "Percentage",
                    f"{percentage:.2f}%",
                )

                remarks = (
                    evaluation.get(
                        "overall_feedback"
                    )
                    or
                    evaluation.get(
                        "remarks"
                    )
                    or
                    "No overall feedback provided."
                )

                st.html(
                    f"""
                    <div class="remark-card">
                        <div class="remark-label">
                            Overall Performance
                        </div>
                        {safe_html(remarks)}
                    </div>
                    """
                )

                criteria = evaluation.get(
                    "criteria",
                    [],
                )

                deductions = evaluation.get(
                    "deductions",
                    [],
                )

                if criteria:

                    st.markdown(
                        "#### Criterion Evaluation"
                    )

                    for criterion in criteria:

                        criterion_name = safe_html(
                            criterion.get(
                                "criterion",
                                criterion.get(
                                    "name",
                                    "Criterion",
                                ),
                            )
                        )

                        score = safe_number(
                            criterion.get(
                                "score",
                                0,
                            )
                        )

                        max_score = safe_number(
                            criterion.get(
                                "max_score",
                                criterion.get(
                                    "max_marks",
                                    0,
                                ),
                            )
                        )

                        evidence = safe_html(
                            criterion.get(
                                "evidence",
                                "",
                            )
                        )

                        feedback = safe_html(
                            criterion.get(
                                "feedback",
                                "",
                            )
                        )

                        st.html(
                            f"""
                            <div class="deduction-card">
                                <div class="deduction-question">
                                    {criterion_name}
                                </div>
                                <div class="deduction-marks">
                                    Score:
                                    {format_score(score)}
                                    /
                                    {format_score(max_score)}
                                </div>
                                <div class="deduction-reason">
                                    <strong>Evidence:</strong>
                                    {evidence or "—"}
                                </div>
                                <div class="deduction-reason">
                                    <strong>Feedback:</strong>
                                    {feedback or "—"}
                                </div>
                            </div>
                            """
                        )

                elif deductions:

                    st.markdown(
                        "#### Mark Deductions"
                    )

                    for deduction in deductions:

                        st.html(
                            f"""
                            <div class="deduction-card">
                                <div class="deduction-question">
                                    {safe_html(
                                        deduction.get(
                                            "question",
                                            "Question",
                                        )
                                    )}
                                </div>
                                <div class="deduction-marks">
                                    Marks Deducted:
                                    {format_score(
                                        deduction.get(
                                            "marks_deducted",
                                            0,
                                        )
                                    )}
                                </div>
                                <div class="deduction-reason">
                                    <strong>Reason:</strong>
                                    {safe_html(
                                        deduction.get(
                                            "reason",
                                            "",
                                        )
                                    )}
                                </div>
                            </div>
                            """
                        )

                per_student_snapshot = {
                    "assignment_file":
                        assignment_snapshot,
                    "rubric_file":
                        rubric_snapshot,
                    "submission_file":
                        submission_snapshot,
                }

                student_report_format = st.selectbox(
                    "Report Format",
                    options=[
                        "PDF",
                        "TXT",
                        "Word (DOCX)",
                    ],
                    index=0,
                    key=(
                        "student_report_format_"
                        f"{index}"
                    ),
                )

                try:

                    (
                        report_bytes,
                        report_file_name,
                        report_mime_type,
                    ) = build_evaluation_report_download(
                        evaluation=evaluation,
                        snapshot=per_student_snapshot,
                        report_format=student_report_format,
                    )

                    st.download_button(
                        label=(
                            "Download Student Report "
                            f"({student_report_format})"
                        ),
                        data=report_bytes,
                        file_name=report_file_name,
                        mime=report_mime_type,
                        key=(
                            "download_student_report_"
                            f"{index}"
                        ),
                        use_container_width=True,
                        on_click="ignore",
                    )

                except RuntimeError as export_error:

                    st.error(
                        str(
                            export_error
                        )
                    )

        # -------------------------------------------------
        # DOWNLOAD ZIP BATCH MARKS XLSX
        # -------------------------------------------------

        if batch_results:

            is_zip_batch = any(
                (
                    item.get(
                        "submission_snapshot",
                        {}
                    ).get(
                        "source_mode"
                    )
                    == "zip"
                )
                for item in batch_results
            )

            if is_zip_batch:

                st.markdown(
                    "### Batch Marks Spreadsheet"
                )

                st.caption(
                    "The Excel summary contains the mandatory Report Generation columns. "
                    "Choose Save to store it in the database for Report Generation, or Download to export the XLSX locally."
                )

                try:

                    batch_xlsx_bytes = (
                        build_batch_marks_xlsx(
                            batch_results=
                                batch_results,
                            assignment_snapshot=
                                assignment_snapshot,
                        )
                    )

                    xlsx_name = batch_marks_xlsx_filename(batch_results)
                    report_rows = build_report_ready_batch_rows(
                        batch_results,
                        assignment_snapshot,
                    )
                    save_token = hashlib.sha256(
                        batch_xlsx_bytes
                    ).hexdigest()

                    save_col, download_col = st.columns(2)

                    with save_col:
                        if st.button(
                            "Save Batch Marks",
                            type="primary",
                            key="save_batch_marks_db",
                            use_container_width=True,
                        ):
                            if (
                                st.session_state.get(
                                    "saved_reporting_batch_token"
                                )
                                == save_token
                            ):
                                st.info(
                                    "This batch is already saved in the database."
                                )
                            else:
                                with st.spinner(
                                    "Saving batch results..."
                                ):
                                    save_result = save_batch_evaluation_results({
                                        "source_filename": assignment_snapshot.get(
                                            "name",
                                            "Assignment",
                                        ),
                                        "xlsx_filename": xlsx_name,
                                        "xlsx_base64": base64.b64encode(
                                            batch_xlsx_bytes
                                        ).decode("ascii"),
                                        "rows": report_rows,
                                    })

                                st.session_state[
                                    "saved_reporting_batch_token"
                                ] = save_token
                                st.session_state[
                                    "saved_reporting_batch_id"
                                ] = save_result.get("batch_id")

                                st.success(
                                    "Batch results saved successfully. "
                                    "They are now available in Report Generation."
                                )

                    with download_col:
                        st.download_button(
                            label="Download Batch Marks (XLSX)",
                            data=batch_xlsx_bytes,
                            file_name=xlsx_name,
                            mime=(
                                "application/vnd.openxmlformats-officedocument."
                                "spreadsheetml.sheet"
                            ),
                            key="download_batch_marks_xlsx",
                            use_container_width=True,
                            on_click="ignore",
                        )

                except Exception as xlsx_error:

                    st.error(
                        "Unable to prepare the XLSX summary. "
                        + str(
                            xlsx_error
                        )
                    )


        # -------------------------------------------------
        # DOWNLOAD ALL REPORTS
        # -------------------------------------------------

        if len(
            successful_results
        ) > 1:

            st.markdown(
                "### Download All Student Reports"
            )

            st.caption(
                "Choose one format. EduAI will create one report "
                "per student and package all reports into a ZIP file."
            )

            batch_report_format = st.selectbox(
                "Batch Report Format",
                options=[
                    "PDF",
                    "TXT",
                    "Word (DOCX)",
                ],
                index=0,
                key="batch_report_format",
            )

            try:

                batch_zip_bytes = (
                    build_batch_reports_zip(
                        batch_results=
                            batch_results,
                        assignment_snapshot=
                            assignment_snapshot,
                        rubric_snapshot=
                            rubric_snapshot,
                        report_format=
                            batch_report_format,
                    )
                )

                st.download_button(
                    label=(
                        "Download All Reports "
                        f"({batch_report_format})"
                    ),
                    data=batch_zip_bytes,
                    file_name=(
                        "EduAI_Student_Evaluation_Reports.zip"
                    ),
                    mime="application/zip",
                    key="download_all_evaluation_reports",
                    use_container_width=True,
                    on_click="ignore",
                )

            except RuntimeError as export_error:

                st.error(
                    str(
                        export_error
                    )
                )


# =========================================================
# OPEN MODAL AFTER EVALUATE CLICK
# =========================================================

if evaluate_clicked:

    try:

        prepared_submission_snapshots = (
            prepare_submission_snapshots(
                submission_files
            )
        )

        st.session_state.ai_evaluation_snapshot = {
            "assignment_file": create_file_snapshot(
                assignment_file
            ),
            "rubric_file": get_active_rubric_snapshot(),
            "submission_files":
                prepared_submission_snapshots,
        }

        st.session_state.ai_evaluation_result = None
        st.session_state.ai_evaluation_batch_results = []
        st.session_state.ai_evaluation_error = None
        st.session_state.ai_evaluation_phase = (
            "evaluating"
        )

        evaluation_modal()

    except Exception as submission_error:

        st.error(
            "Unable to prepare the student submissions. "
            + str(
                submission_error
            )
        )