import streamlit as st

import sys
import os
import io
import html
import json

from docx import Document
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


# =========================================================
# PROJECT ROOT
# =========================================================

PROJECT_ROOT = os.path.abspath(
    os.path.join(
        os.path.dirname(__file__),
        ".."
    )
)

if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)


# =========================================================
# PROJECT IMPORTS
# =========================================================

from api_client import (
    generate_questions,
    generate_questions_from_file,
    save_question_set,
    get_question_sets,
    get_question_set,
    update_question,
    delete_question_set,
    download_pdf,
    download_docx,
)

from backend.services.pdf_generator import (
    generate_question_paper_pdf,
)


# =========================================================
# PAGE CONFIG
# =========================================================

st.set_page_config(
    page_title="Question Generation",
    layout="wide",
    initial_sidebar_state="collapsed",
)


# =========================================================
# SESSION STATE
# =========================================================

DEFAULT_STATE = {
    "generated_questions": None,
    "generation_modal_phase": "idle",
    "generation_error": None,
    "generation_snapshot": None,
    "preview_pdf": None,
    "save_success_message": None,

    # Question generation logs
    "selected_log_set": None,
    "selected_log_set_id": None,
    "confirm_delete_log_id": None,
    "question_log_mode": "view",
    "editing_question_id": None,
    "question_set_total_overrides": {},
    "logs_flash_message": None,
}


for key, value in DEFAULT_STATE.items():

    if key not in st.session_state:
        st.session_state[key] = value


# =========================================================
# MEMORY UPLOADED FILE
# =========================================================

class MemoryUploadedFile(io.BytesIO):

    def __init__(
        self,
        data,
        name="uploaded_file",
        mime_type="application/octet-stream",
    ):

        super().__init__(data)

        self.name = name
        self.type = mime_type
        self.size = len(data)


# =========================================================
# FILE SNAPSHOT
# =========================================================

def create_file_snapshot(uploaded_file):

    if uploaded_file is None:
        return None

    try:
        file_bytes = uploaded_file.getvalue()

    except Exception:

        try:
            uploaded_file.seek(0)
        except Exception:
            pass

        file_bytes = uploaded_file.read()

    return {
        "bytes": file_bytes,
        "name": getattr(
            uploaded_file,
            "name",
            "uploaded_file",
        ),
        "type": getattr(
            uploaded_file,
            "type",
            "application/octet-stream",
        ),
    }


# =========================================================
# REBUILD FILE
# =========================================================

def rebuild_file(file_snapshot):

    if not file_snapshot:
        return None

    return MemoryUploadedFile(
        data=file_snapshot["bytes"],
        name=file_snapshot.get(
            "name",
            "uploaded_file",
        ),
        mime_type=file_snapshot.get(
            "type",
            "application/octet-stream",
        ),
    )


# =========================================================
# PDF OUTPUT -> BYTES
# =========================================================

def get_pdf_bytes(pdf_file):

    if isinstance(pdf_file, bytes):
        return pdf_file

    if isinstance(pdf_file, bytearray):
        return bytes(pdf_file)

    if isinstance(pdf_file, str):

        if os.path.exists(pdf_file):

            with open(
                pdf_file,
                "rb",
            ) as file:

                return file.read()

    if hasattr(
        pdf_file,
        "getvalue",
    ):
        return pdf_file.getvalue()

    if hasattr(
        pdf_file,
        "read",
    ):

        try:
            pdf_file.seek(0)
        except Exception:
            pass

        return pdf_file.read()

    raise TypeError(
        "PDF generator returned an unsupported file type."
    )



# =========================================================
# DOWNLOAD RESPONSE -> BYTES
# =========================================================

def get_download_bytes(download_data):

    if download_data is None:
        return None

    if isinstance(download_data, bytes):
        return download_data

    if isinstance(download_data, bytearray):
        return bytes(download_data)

    # Supports requests.Response if api_client returns it directly
    if hasattr(download_data, "content"):

        content = download_data.content

        if isinstance(content, bytes):
            return content

        if isinstance(content, bytearray):
            return bytes(content)

    if hasattr(download_data, "getvalue"):
        return download_data.getvalue()

    if hasattr(download_data, "read"):

        try:
            download_data.seek(0)
        except Exception:
            pass

        return download_data.read()

    raise TypeError(
        "Download API returned an unsupported file type."
    )



# =========================================================
# NORMALIZE QUESTION OPTIONS
# =========================================================

def normalize_options(options):

    if options is None:
        return []

    if isinstance(options, list):
        return [str(item) for item in options]

    if isinstance(options, tuple):
        return [str(item) for item in options]

    if isinstance(options, str):

        value = options.strip()

        if not value:
            return []

        try:
            parsed = json.loads(value)

            if isinstance(parsed, list):
                return [str(item) for item in parsed]

        except Exception:
            pass

        return [value]

    return [str(options)]


# =========================================================
# BUILD PDF FROM SELECTED SAVED QUESTION SET
# Fallback used when backend export endpoint is unavailable.
# =========================================================

def build_question_set_pdf(selected):

    buffer = io.BytesIO()

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=17 * mm,
        bottomMargin=17 * mm,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "EduAITitle",
        parent=styles["Heading1"],
        alignment=TA_CENTER,
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#0F172A"),
        spaceAfter=5,
    )

    subtitle_style = ParagraphStyle(
        "EduAISubtitle",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#64748B"),
        spaceAfter=14,
    )

    question_style = ParagraphStyle(
        "EduAIQuestion",
        parent=styles["BodyText"],
        fontSize=10.5,
        leading=15,
        textColor=colors.HexColor("#1E293B"),
        spaceAfter=5,
    )

    meta_style = ParagraphStyle(
        "EduAIMeta",
        parent=styles["BodyText"],
        fontSize=8.5,
        leading=12,
        textColor=colors.HexColor("#64748B"),
        spaceAfter=9,
    )

    title = str(
        selected.get(
            "title",
            "EduAI Question Paper",
        )
    )

    topic = str(
        selected.get(
            "topic",
            "-",
        )
    )

    difficulty = str(
        selected.get(
            "difficulty",
            "-",
        )
    )

    total_marks = selected.get(
        "total_marks",
        0,
    )

    questions = selected.get(
        "questions",
        [],
    )

    story = [
        Paragraph(
            "EduAI",
            title_style,
        ),
        Paragraph(
            html.escape(title),
            title_style,
        ),
        Paragraph(
            (
                f"Topic: {html.escape(topic)}"
                f" &nbsp;&nbsp; | &nbsp;&nbsp; "
                f"Difficulty: {html.escape(difficulty)}"
                f" &nbsp;&nbsp; | &nbsp;&nbsp; "
                f"Total Marks: {html.escape(str(total_marks))}"
            ),
            subtitle_style,
        ),
        Spacer(1, 4),
    ]

    for index, question in enumerate(
        questions,
        start=1,
    ):

        number = question.get(
            "question_number",
            index,
        )

        question_text = question.get(
            "question_text",
            question.get(
                "question",
                "",
            ),
        )

        question_type = question.get(
            "type",
            "-",
        )

        marks = question.get(
            "marks",
            0,
        )

        story.append(
            Paragraph(
                (
                    f"<b>Question {html.escape(str(number))}</b><br/>"
                    f"{html.escape(str(question_text))}"
                ),
                question_style,
            )
        )

        for option in normalize_options(
            question.get(
                "options"
            )
        ):
            story.append(
                Paragraph(
                    f"&nbsp;&nbsp;&nbsp;&nbsp;{html.escape(option)}",
                    styles["BodyText"],
                )
            )

        story.append(
            Paragraph(
                (
                    f"Type: {html.escape(str(question_type))}"
                    f" &nbsp;&nbsp; | &nbsp;&nbsp; "
                    f"Marks: {html.escape(str(marks))}"
                ),
                meta_style,
            )
        )

        story.append(
            Spacer(
                1,
                4,
            )
        )

    document.build(
        story
    )

    buffer.seek(0)

    return buffer.getvalue()


# =========================================================
# BUILD DOCX FROM SELECTED SAVED QUESTION SET
# Fallback used when backend export endpoint is unavailable.
# =========================================================

def build_question_set_docx(selected):

    document = Document()

    heading = document.add_heading(
        "EduAI",
        level=0,
    )

    heading.alignment = 1

    title = str(
        selected.get(
            "title",
            "EduAI Question Paper",
        )
    )

    title_paragraph = document.add_paragraph()

    title_paragraph.alignment = 1

    title_run = title_paragraph.add_run(
        title
    )

    title_run.bold = True

    topic = str(
        selected.get(
            "topic",
            "-",
        )
    )

    difficulty = str(
        selected.get(
            "difficulty",
            "-",
        )
    )

    total_marks = selected.get(
        "total_marks",
        0,
    )

    meta = document.add_paragraph()

    meta.alignment = 1

    meta.add_run(
        f"Topic: {topic}   |   "
        f"Difficulty: {difficulty}   |   "
        f"Total Marks: {total_marks}"
    )

    document.add_paragraph("")

    questions = selected.get(
        "questions",
        [],
    )

    for index, question in enumerate(
        questions,
        start=1,
    ):

        number = question.get(
            "question_number",
            index,
        )

        question_text = question.get(
            "question_text",
            question.get(
                "question",
                "",
            ),
        )

        question_type = question.get(
            "type",
            "-",
        )

        marks = question.get(
            "marks",
            0,
        )

        paragraph = document.add_paragraph()

        number_run = paragraph.add_run(
            f"Question {number}\n"
        )

        number_run.bold = True

        paragraph.add_run(
            str(question_text)
        )

        for option in normalize_options(
            question.get(
                "options"
            )
        ):

            option_paragraph = document.add_paragraph(
                style="List Bullet"
            )

            option_paragraph.add_run(
                option
            )

        meta_paragraph = document.add_paragraph()

        meta_run = meta_paragraph.add_run(
            f"Type: {question_type}   |   Marks: {marks}"
        )

        meta_run.italic = True

    buffer = io.BytesIO()

    document.save(
        buffer
    )

    buffer.seek(0)

    return buffer.getvalue()


# =========================================================
# RELIABLE PDF EXPORT
# Uses backend export first, then falls back to local PDF.
# =========================================================

def get_reliable_pdf_export(
    selected_id,
    selected,
):

    try:

        pdf_data = get_download_bytes(
            download_pdf(
                selected_id
            )
        )

        if (
            pdf_data
            and pdf_data.startswith(
                b"%PDF"
            )
        ):

            return pdf_data

    except Exception as exc:

        print(
            "Backend PDF export fallback:",
            exc,
        )

    return build_question_set_pdf(
        selected
    )


# =========================================================
# RELIABLE DOCX EXPORT
# Uses backend export first, then falls back to local DOCX.
# DOCX files are ZIP containers, so valid data starts with PK.
# =========================================================

def get_reliable_docx_export(
    selected_id,
    selected,
):

    try:

        docx_data = get_download_bytes(
            download_docx(
                selected_id
            )
        )

        if (
            docx_data
            and docx_data.startswith(
                b"PK"
            )
        ):

            return docx_data

    except Exception as exc:

        print(
            "Backend DOCX export fallback:",
            exc,
        )

    return build_question_set_docx(
        selected
    )



# =========================================================
# CALCULATE TOTAL MARKS FROM QUESTIONS
# =========================================================

def calculate_total_marks_from_questions(
    questions,
):

    total = 0

    for question in (
        questions
        or []
    ):

        try:

            total += int(
                float(
                    question.get(
                        "marks",
                        0,
                    )
                    or 0
                )
            )

        except (
            TypeError,
            ValueError,
        ):

            continue

    return total


# =========================================================
# NORMALIZE QUESTION SET DETAILS
# =========================================================

def normalize_question_set_details(
    data,
    fallback_id=None,
):

    if not isinstance(
        data,
        dict,
    ):

        return {}


    # -----------------------------------------------------
    # Backend shape:
    #
    # {
    #     "question_set": {...},
    #     "questions": [...]
    # }
    # -----------------------------------------------------

    if isinstance(
        data.get(
            "question_set"
        ),
        dict,
    ):

        normalized = (
            data[
                "question_set"
            ].copy()
        )


        if "questions" in data:

            normalized[
                "questions"
            ] = data[
                "questions"
            ]


    # -----------------------------------------------------
    # Already-flat response
    # -----------------------------------------------------

    else:

        normalized = (
            data.copy()
        )


    # -----------------------------------------------------
    # GUARANTEE QUESTION SET ID
    # -----------------------------------------------------

    existing_id = (
        normalized.get(
            "id"
        )
        or
        normalized.get(
            "set_id"
        )
        or
        fallback_id
    )


    if existing_id is not None:

        normalized[
            "id"
        ] = existing_id


    # -----------------------------------------------------
    # KEEP TOTAL MARKS CONSISTENT WITH INDIVIDUAL QUESTIONS
    # -----------------------------------------------------

    normalized_questions = normalized.get(
        "questions",
        [],
    )


    if normalized_questions:

        normalized[
            "total_marks"
        ] = calculate_total_marks_from_questions(
            normalized_questions
        )


    return normalized


# =========================================================
# CSS
# =========================================================

st.html(
"""
<style>

/* =====================================================
   REMOVE STREAMLIT DEFAULT UI
===================================================== */

[data-testid="stSidebar"] {
    display: none !important;
}

[data-testid="collapsedControl"] {
    display: none !important;
}

header {
    display: none !important;
}

footer {
    display: none !important;
}

#MainMenu {
    visibility: hidden !important;
}


/* =====================================================
   PAGE BACKGROUND
===================================================== */

.stApp {

    background:
        radial-gradient(
            circle at 8% 12%,
            rgba(99, 102, 241, 0.08),
            transparent 28%
        ),
        radial-gradient(
            circle at 92% 82%,
            rgba(59, 130, 246, 0.06),
            transparent 30%
        ),
        #f8fafc;
}


/* =====================================================
   MAIN CONTAINER
===================================================== */

.block-container {

    max-width: 1180px !important;

    padding-top: 45px !important;
    padding-bottom: 80px !important;

    padding-left: 28px !important;
    padding-right: 28px !important;
}



/* =====================================================
   BACK BUTTON
   Uses a normal link to the app root because Dashboard
   is a callable default page in st.navigation.
===================================================== */

.back-link-wrap {

    padding-top:
        31px;

    display:
        flex;

    justify-content:
        flex-start;
}


.back-link {

    width:
        40px;

    height:
        40px;

    display:
        inline-flex;

    align-items:
        center;

    justify-content:
        center;

    background:
        #ffffff;

    color:
        #334155 !important;

    border:
        1px solid #dbe3ed;

    border-radius:
        12px;

    font-size:
        24px;

    font-weight:
        500;

    line-height:
        1;

    text-decoration:
        none !important;

    box-shadow:
        0 4px 12px
        rgba(15, 23, 42, 0.06);

    transition:
        all 0.18s ease;
}


.back-link:hover {

    background:
        #f8fafc;

    color:
        #4f46e5 !important;

    border-color:
        #c7d2fe;

    transform:
        translateX(-1px);

    box-shadow:
        0 6px 16px
        rgba(15, 23, 42, 0.08);
}


.back-link:focus {

    outline:
        none;

    box-shadow:
        0 0 0 3px
        rgba(79, 70, 229, 0.10);
}


/* =====================================================
   PAGE HEADER
===================================================== */

.page-header {
    margin-bottom: 32px;
}


.page-label {

    color: #4f46e5;

    font-size: 12px;

    font-weight: 750;

    letter-spacing: 1.2px;

    text-transform: uppercase;

    margin-bottom: 12px;
}


.page-title {

    color: #0f172a;

    font-size: 38px;

    font-weight: 780;

    letter-spacing: -1.1px;

    line-height: 1.15;

    margin-bottom: 12px;
}


.page-description {

    color: #64748b;

    font-size: 15px;

    line-height: 1.7;

    max-width: 750px;
}


/* =====================================================
   SECTION CONTAINERS
===================================================== */

[data-testid="stVerticalBlockBorderWrapper"] {

    background:
        rgba(255, 255, 255, 0.97);

    border:
        1px solid #e2e8f0 !important;

    border-radius:
        20px !important;

    box-shadow:
        0 8px 28px
        rgba(15, 23, 42, 0.045);

    margin-bottom: 20px;

    transition:
        border-color 0.2s ease,
        box-shadow 0.2s ease;
}


[data-testid="stVerticalBlockBorderWrapper"]:hover {

    border-color:
        #d3d9e7 !important;

    box-shadow:
        0 12px 32px
        rgba(15, 23, 42, 0.06);
}


/* =====================================================
   SECTION HEADER
===================================================== */

.section-header {

    padding:
        3px
        1px
        17px
        1px;
}


.section-number {

    color: #4f46e5;

    font-size: 11px;

    font-weight: 750;

    letter-spacing: 1px;

    text-transform: uppercase;

    margin-bottom: 8px;
}


.section-title {

    color: #0f172a;

    font-size: 19px;

    font-weight: 760;

    letter-spacing: -0.3px;

    margin-bottom: 6px;
}


.section-description {

    color: #64748b;

    font-size: 13.5px;

    line-height: 1.6;
}


/* =====================================================
   LABELS
===================================================== */

label,
[data-testid="stWidgetLabel"] {

    color:
        #334155 !important;

    font-size:
        13.5px !important;

    font-weight:
        650 !important;
}


/* =====================================================
   TEXT INPUT
===================================================== */

[data-testid="stTextInput"] input {

    min-height:
        46px !important;

    border-radius:
        11px !important;

    border:
        1px solid #dce3ed !important;

    background:
        #ffffff !important;

    color:
        #0f172a !important;

    font-size:
        14px !important;

    box-shadow:
        none !important;
}


[data-testid="stTextInput"] input:focus {

    border-color:
        #6366f1 !important;

    box-shadow:
        0 0 0 3px
        rgba(99, 102, 241, 0.10) !important;
}


/* =====================================================
   NUMBER INPUT
===================================================== */

[data-testid="stNumberInput"] input {

    min-height:
        46px !important;

    background:
        #ffffff !important;

    color:
        #0f172a !important;

    border-color:
        #dce3ed !important;
}


/* =====================================================
   SELECTBOX
===================================================== */

div[data-baseweb="select"] > div {

    min-height:
        46px !important;

    border-radius:
        11px !important;

    border-color:
        #dce3ed !important;

    background:
        #ffffff !important;

    box-shadow:
        none !important;
}


div[data-baseweb="select"] > div:hover {

    border-color:
        #a5b4fc !important;
}


/* =====================================================
   RADIO BUTTONS
===================================================== */

[data-testid="stRadio"] > div {

    gap:
        10px !important;
}


[data-testid="stRadio"] label {

    background:
        #f8fafc;

    border:
        1px solid #e2e8f0;

    border-radius:
        10px;

    padding:
        9px 14px;

    transition:
        all 0.18s ease;
}


[data-testid="stRadio"] label:hover {

    border-color:
        #a5b4fc;

    background:
        #f5f7ff;
}


/* =====================================================
   FILE UPLOADER
===================================================== */

[data-testid="stFileUploaderDropzone"] {

    background:
        #f8fafc !important;

    border:
        1px dashed #cbd5e1 !important;

    border-radius:
        14px !important;

    padding:
        23px !important;

    transition:
        all 0.20s ease;
}


[data-testid="stFileUploaderDropzone"]:hover {

    background:
        #f6f7ff !important;

    border-color:
        #818cf8 !important;
}


/* =====================================================
   TEXT AREA
===================================================== */

[data-testid="stTextArea"] textarea {

    border-radius:
        12px !important;

    border:
        1px solid #dce3ed !important;

    background:
        #ffffff !important;

    color:
        #0f172a !important;

    font-size:
        14px !important;

    line-height:
        1.6 !important;

    box-shadow:
        none !important;
}


[data-testid="stTextArea"] textarea:focus {

    border-color:
        #6366f1 !important;

    box-shadow:
        0 0 0 3px
        rgba(99, 102, 241, 0.10) !important;
}


/* =====================================================
   MAIN BUTTONS
===================================================== */

.stButton > button {

    min-height:
        47px !important;

    width:
        100% !important;

    border-radius:
        11px !important;

    border:
        1px solid #4f46e5 !important;

    background:
        #4f46e5 !important;

    color:
        #ffffff !important;

    font-size:
        13.5px !important;

    font-weight:
        700 !important;

    box-shadow:
        0 6px 16px
        rgba(79, 70, 229, 0.15) !important;

    transition:
        all 0.20s ease !important;
}


.stButton > button:hover {

    background:
        #4338ca !important;

    border-color:
        #4338ca !important;

    color:
        #ffffff !important;

    transform:
        translateY(-1px);

    box-shadow:
        0 9px 21px
        rgba(67, 56, 202, 0.20) !important;
}


/* =====================================================
   DOWNLOAD BUTTON
===================================================== */

.stDownloadButton > button {

    min-height:
        47px !important;

    width:
        100% !important;

    border-radius:
        11px !important;

    border:
        1px solid #4f46e5 !important;

    background:
        #ffffff !important;

    color:
        #4f46e5 !important;

    font-size:
        13.5px !important;

    font-weight:
        700 !important;

    transition:
        all 0.20s ease !important;
}


.stDownloadButton > button:hover {

    background:
        #4f46e5 !important;

    color:
        #ffffff !important;
}


/* =====================================================
   CURRENT CONFIGURATION
===================================================== */

.generation-summary {

    background:
        #f8fafc;

    border:
        1px solid #e2e8f0;

    border-radius:
        13px;

    padding:
        18px 20px;

    margin-bottom:
        17px;
}


.summary-title {

    color:
        #0f172a;

    font-size:
        14px;

    font-weight:
        700;

    margin-bottom:
        6px;
}


.summary-text {

    color:
        #64748b;

    font-size:
        13px;

    line-height:
        1.7;
}


/* =====================================================
   MODAL FRAME
   Compact desktop dialog; content scrolls inside
===================================================== */

div[data-testid="stDialog"] {

    padding:
        12px !important;

    overflow:
        hidden !important;
}


div[data-testid="stDialog"] div[role="dialog"] {

    width:
        min(1120px, 92vw) !important;

    max-width:
        92vw !important;

    height:
        auto !important;

    max-height:
        86vh !important;

    margin:
        auto !important;

    border-radius:
        18px !important;

    overflow:
        hidden !important;

    box-shadow:
        0 24px 70px
        rgba(15, 23, 42, 0.20) !important;
}


/* =====================================================
   QUESTION LOGS SCROLL AREA
   This is the main fix: the modal stays compact and
   View/Edit content scrolls inside the popup.
===================================================== */

.st-key-question_logs_scroll {

    height:
        60vh !important;

    max-height:
        60vh !important;

    overflow-y:
        auto !important;

    overflow-x:
        hidden !important;

    padding-right:
        8px !important;

    scrollbar-width:
        thin;

    scrollbar-color:
        #cbd5e1 transparent;
}


.st-key-question_logs_scroll::-webkit-scrollbar {

    width:
        8px;
}


.st-key-question_logs_scroll::-webkit-scrollbar-track {

    background:
        transparent;
}


.st-key-question_logs_scroll::-webkit-scrollbar-thumb {

    background:
        #cbd5e1;

    border-radius:
        999px;
}


.st-key-question_logs_scroll::-webkit-scrollbar-thumb:hover {

    background:
        #94a3b8;
}


/* =====================================================
   GENERATION MODAL SCROLL AREA
===================================================== */

.st-key-generation_modal_scroll {

    max-height:
        76vh !important;

    overflow-y:
        auto !important;

    overflow-x:
        hidden !important;

    padding-right:
        8px !important;

    scrollbar-width:
        thin;

    scrollbar-color:
        #cbd5e1 transparent;
}


.st-key-generation_modal_scroll::-webkit-scrollbar {

    width:
        8px;
}


.st-key-generation_modal_scroll::-webkit-scrollbar-thumb {

    background:
        #cbd5e1;

    border-radius:
        999px;
}


/* Compact widget spacing inside dialogs */

div[role="dialog"] [data-testid="stVerticalBlock"] {

    gap:
        0.60rem !important;
}


div[role="dialog"] .logs-header,
div[role="dialog"] .modal-header {

    margin-bottom:
        8px !important;
}


div[role="dialog"] .stButton > button,
div[role="dialog"] .stDownloadButton > button {

    min-height:
        40px !important;
}


/* Smaller laptop screens */

@media (max-height: 800px) {

    .st-key-question_logs_scroll {

        height:
            56vh !important;

        max-height:
            56vh !important;
    }


    .st-key-generation_modal_scroll {

        max-height:
            70vh !important;
    }
}


/* Narrow screens */

@media (max-width: 850px) {

    div[data-testid="stDialog"] {

        padding:
            6px !important;
    }


    div[data-testid="stDialog"] div[role="dialog"] {

        width:
            96vw !important;

        max-width:
            96vw !important;

        max-height:
            90vh !important;

        border-radius:
            14px !important;
    }


    .st-key-question_logs_scroll {

        height:
            72vh !important;

        max-height:
            72vh !important;
    }
}


/* =====================================================
   MODAL HEADER
===================================================== */

.modal-header {

    padding:
        3px
        3px
        8px
        3px;
}


.modal-title {

    color:
        #0f172a;

    font-size:
        23px;

    font-weight:
        780;

    letter-spacing:
        -0.5px;

    margin-bottom:
        7px;
}


.modal-description {

    color:
        #64748b;

    font-size:
        13.5px;

    line-height:
        1.65;

    max-width:
        750px;
}


/* =====================================================
   MODAL CONFIGURATION
===================================================== */

.modal-config {

    background:
        #f8fafc;

    border:
        1px solid #e2e8f0;

    border-radius:
        13px;

    padding:
        14px 16px;

    margin:
        7px
        0
        5px
        0;

    color:
        #64748b;

    font-size:
        12.5px;

    line-height:
        1.7;
}


.modal-config strong {

    color:
        #334155;

    font-weight:
        700;
}


/* =====================================================
   LOADING PANEL
===================================================== */

.loading-card {

    background:
        linear-gradient(
            135deg,
            #ffffff,
            #f7f8ff
        );

    border:
        1px solid #c7d2fe;

    border-radius:
        17px;

    padding:
        27px;

    margin:
        10px
        0;

    box-shadow:
        0 10px 30px
        rgba(79, 70, 229, 0.08);
}


.loading-title {

    color:
        #0f172a;

    font-size:
        18px;

    font-weight:
        760;

    margin-bottom:
        8px;
}


.loading-text {

    color:
        #64748b;

    font-size:
        13.5px;

    line-height:
        1.65;

    margin-bottom:
        21px;
}


.loading-track {

    position:
        relative;

    width:
        100%;

    height:
        7px;

    background:
        #e8eaf5;

    border-radius:
        999px;

    overflow:
        hidden;
}


.loading-bar {

    position:
        absolute;

    top:
        0;

    left:
        -40%;

    width:
        38%;

    height:
        100%;

    background:
        linear-gradient(
            90deg,
            #4f46e5,
            #6366f1,
            #818cf8
        );

    border-radius:
        999px;

    animation:
        modal-loading
        1.2s
        ease-in-out
        infinite;
}


@keyframes modal-loading {

    0% {
        left: -40%;
    }

    50% {
        left: 45%;
    }

    100% {
        left: 105%;
    }
}


.loading-note {

    color:
        #94a3b8;

    font-size:
        12px;

    margin-top:
        13px;
}


/* =====================================================
   SUCCESS
===================================================== */

.success-card {

    background:
        #ffffff;

    border:
        1px solid #d8e0ea;

    border-radius:
        16px;

    padding:
        22px 23px;

    margin:
        10px
        0
        16px
        0;

    box-shadow:
        0 7px 20px
        rgba(15, 23, 42, 0.04);
}


.success-title {

    color:
        #0f172a;

    font-size:
        16px;

    font-weight:
        760;

    margin-bottom:
        7px;
}


.success-text {

    color:
        #64748b;

    font-size:
        13px;

    line-height:
        1.65;
}


/* =====================================================
   ERROR
===================================================== */

.error-card {

    background:
        #fffafa;

    border:
        1px solid #fecaca;

    border-radius:
        16px;

    padding:
        21px 22px;

    margin:
        10px
        0;
}


.error-title {

    color:
        #991b1b;

    font-size:
        15px;

    font-weight:
        750;

    margin-bottom:
        6px;
}


.error-text {

    color:
        #7f1d1d;

    font-size:
        13px;

    line-height:
        1.6;
}


/* =====================================================
   PDF PREVIEW AREA
===================================================== */

.pdf-preview-label {

    color:
        #0f172a;

    font-size:
        14px;

    font-weight:
        700;

    margin:
        4px
        0
        8px
        0;
}


.pdf-preview-note {

    background:
        #f8fafc;

    border:
        1px solid #e2e8f0;

    border-radius:
        12px;

    color:
        #64748b;

    font-size:
        12.5px;

    line-height:
        1.6;

    padding:
        12px
        14px;

    margin-bottom:
        8px;
}


/* =====================================================
   FINAL ACTION LABEL
===================================================== */

.modal-actions-label {

    color:
        #64748b;

    font-size:
        12.5px;

    margin-top:
        6px;

    margin-bottom:
        3px;
}


/* =====================================================
   ALERT
===================================================== */

[data-testid="stAlert"] {

    border-radius:
        12px !important;
}


[data-testid="stAlert"] svg {

    display:
        none !important;
}



/* =====================================================
   VIEW LOGS BUTTON
===================================================== */

.st-key-view_logs_button button {

    width: 100% !important;
    min-height: 40px !important;
    height: 40px !important;

    background: #ffffff !important;
    color: #4f46e5 !important;

    border: 1px solid #c7d2fe !important;
    border-radius: 10px !important;

    font-size: 12.5px !important;
    font-weight: 700 !important;

    box-shadow:
        0 3px 10px
        rgba(15, 23, 42, 0.05) !important;

    visibility: visible !important;
    opacity: 1 !important;

    transition: all 0.20s ease !important;
}


.st-key-view_logs_button button:hover {

    background: #4f46e5 !important;
    color: #ffffff !important;

    border-color: #4f46e5 !important;

    transform: translateY(-1px);

    box-shadow:
        0 7px 16px
        rgba(79, 70, 229, 0.16) !important;
}


/* =====================================================
   LOGS MODAL
===================================================== */

.logs-header {
    margin-bottom: 16px;
}


.logs-title {

    color: #0f172a;

    font-size: 23px;
    font-weight: 780;

    letter-spacing: -0.45px;

    margin-bottom: 6px;
}


.logs-description {

    color: #64748b;

    font-size: 13.5px;
    line-height: 1.65;

    max-width: 760px;
}


.logs-summary {

    background: #f8fafc;

    border: 1px solid #e2e8f0;
    border-radius: 13px;

    padding: 14px 16px;

    margin: 8px 0 16px 0;
}


.logs-summary-title {

    color: #0f172a;

    font-size: 13px;
    font-weight: 700;

    margin-bottom: 3px;
}


.logs-summary-text {

    color: #64748b;

    font-size: 12.5px;
    line-height: 1.6;
}


.log-table-header {

    background: #f1f5f9;

    color: #475569;

    padding: 10px 8px;

    border-radius: 8px;

    font-size: 10.5px;
    font-weight: 750;

    text-transform: uppercase;
    letter-spacing: 0.35px;

    min-height: 37px;

    box-sizing: border-box;
}


.log-row-value {

    color: #334155;

    font-size: 12px;
    line-height: 1.5;

    padding: 11px 4px 4px 4px;

    overflow: hidden;

    text-overflow: ellipsis;
    white-space: nowrap;
}


.log-divider {

    border-top: 1px solid #edf1f5;

    margin: 7px 0;
}


.log-badge {

    display: inline-flex;

    align-items: center;

    padding: 5px 9px;

    border-radius: 7px;

    font-size: 11px;
    font-weight: 700;

    margin-top: 7px;
}


.log-easy {

    background: #ecfdf5;
    color: #047857;

    border: 1px solid #a7f3d0;
}


.log-medium {

    background: #fffbeb;
    color: #b45309;

    border: 1px solid #fde68a;
}


.log-hard {

    background: #fef2f2;
    color: #b91c1c;

    border: 1px solid #fecaca;
}


.log-detail {

    background: #ffffff;

    border: 1px solid #dbe3ed;
    border-radius: 15px;

    padding: 20px 21px;

    margin: 20px 0 14px 0;

    box-shadow:
        0 4px 14px
        rgba(15, 23, 42, 0.03);
}


.log-detail-title {

    color: #0f172a;

    font-size: 18px;
    font-weight: 760;

    margin-bottom: 6px;
}


.log-detail-meta {

    color: #64748b;

    font-size: 12.5px;
    line-height: 1.6;
}


.log-question {

    background: #f8fafc;

    border: 1px solid #e2e8f0;
    border-radius: 12px;

    padding: 15px 16px;

    margin-bottom: 10px;
}


.log-question-number {

    color: #4f46e5;

    font-size: 11px;
    font-weight: 750;

    text-transform: uppercase;
    letter-spacing: 0.5px;

    margin-bottom: 7px;
}


.log-question-text {

    color: #334155;

    font-size: 13px;
    line-height: 1.65;
}


.log-question-meta {

    color: #64748b;

    font-size: 11.5px;

    margin-top: 9px;
}


[class*="st-key-log_view_"] button {

    min-height: 36px !important;
    height: 36px !important;

    background: #ffffff !important;
    color: #4f46e5 !important;

    border: 1px solid #c7d2fe !important;
    border-radius: 8px !important;

    font-size: 11px !important;
    font-weight: 700 !important;

    box-shadow: none !important;
}


[class*="st-key-log_view_"] button:hover {

    background: #f5f7ff !important;
    border-color: #818cf8 !important;
    color: #4338ca !important;
}


[class*="st-key-log_delete_"] button,
[class*="st-key-log_confirm_delete_"] button {

    background: #ffffff !important;
    color: #b91c1c !important;

    border: 1px solid #fecaca !important;

    box-shadow: none !important;
}


[class*="st-key-log_delete_"] button:hover,
[class*="st-key-log_confirm_delete_"] button:hover {

    background: #fef2f2 !important;
    border-color: #f87171 !important;
    color: #991b1b !important;
}


[class*="st-key-log_cancel_delete_"] button {

    background: #ffffff !important;
    color: #475569 !important;

    border: 1px solid #cbd5e1 !important;

    box-shadow: none !important;
}


/* =====================================================
   LOG ACTION BUTTON GROUP
===================================================== */

.log-action-wrap [data-testid="stHorizontalBlock"] {
    gap: 6px !important;
}


/* =====================================================
   EDIT BUTTON
===================================================== */

[class*="st-key-log_edit_"] button {

    min-height: 36px !important;
    height: 36px !important;

    background: #4f46e5 !important;
    color: #ffffff !important;

    border: 1px solid #4f46e5 !important;
    border-radius: 8px !important;

    font-size: 11px !important;
    font-weight: 700 !important;

    box-shadow: none !important;
}


[class*="st-key-log_edit_"] button:hover {

    background: #4338ca !important;
    border-color: #4338ca !important;
    color: #ffffff !important;
}


/* =====================================================
   EDIT PANEL
===================================================== */

.edit-panel-header {

    background:
        linear-gradient(
            135deg,
            #ffffff,
            #f8faff
        );

    border: 1px solid #c7d2fe;
    border-radius: 15px;

    padding: 18px 20px;

    margin: 20px 0 14px 0;
}


.edit-panel-title {

    color: #0f172a;

    font-size: 18px;
    font-weight: 760;

    margin-bottom: 6px;
}


.edit-panel-text {

    color: #64748b;

    font-size: 12.5px;
    line-height: 1.6;
}


.edit-question-card {

    background: #ffffff;

    border: 1px solid #e2e8f0;
    border-radius: 14px;

    padding: 16px 17px;

    margin-bottom: 12px;
}


.edit-question-label {

    color: #4f46e5;

    font-size: 11px;
    font-weight: 750;

    text-transform: uppercase;
    letter-spacing: 0.5px;

    margin-bottom: 7px;
}


/* =====================================================
   SAVE EDIT BUTTON
===================================================== */

[class*="st-key-save_question_edit_"] button {

    background: #4f46e5 !important;
    color: #ffffff !important;

    border: 1px solid #4f46e5 !important;

    box-shadow: none !important;
}


[class*="st-key-save_question_edit_"] button:hover {

    background: #4338ca !important;
    border-color: #4338ca !important;
    color: #ffffff !important;
}


/* =====================================================
   BACK TO VIEW BUTTON
===================================================== */

.st-key-back_to_log_view button {

    background: #ffffff !important;
    color: #4f46e5 !important;

    border: 1px solid #c7d2fe !important;

    box-shadow: none !important;
}


.st-key-back_to_log_view button:hover {

    background: #f5f7ff !important;
    border-color: #818cf8 !important;
    color: #4338ca !important;
}



/* =====================================================
   VIEW / EDIT EXPANDER
===================================================== */

div[role="dialog"] [data-testid="stExpander"] {

    border:
        1px solid #dbe3ed !important;

    border-radius:
        14px !important;

    background:
        #ffffff !important;

    box-shadow:
        0 5px 18px
        rgba(15, 23, 42, 0.035) !important;

    overflow:
        hidden !important;

    margin-top:
        14px !important;
}


div[role="dialog"] [data-testid="stExpander"] details {

    border:
        none !important;
}


div[role="dialog"] [data-testid="stExpander"] summary {

    min-height:
        48px !important;

    padding:
        8px 12px !important;

    background:
        #f8fafc !important;

    color:
        #0f172a !important;

    font-size:
        13px !important;

    font-weight:
        750 !important;
}


div[role="dialog"] [data-testid="stExpander"] summary:hover {

    background:
        #f5f7ff !important;
}


[class*="st-key-close_log_view_"] button,
[class*="st-key-close_log_edit_"] button {

    background:
        #ffffff !important;

    color:
        #475569 !important;

    border:
        1px solid #cbd5e1 !important;

    box-shadow:
        none !important;
}


[class*="st-key-close_log_view_"] button:hover,
[class*="st-key-close_log_edit_"] button:hover {

    background:
        #f8fafc !important;

    color:
        #0f172a !important;

    border-color:
        #94a3b8 !important;
}


/* =====================================================
   MOBILE
===================================================== */

@media (max-width: 850px) {

    .block-container {

        padding-top:
            28px !important;

        padding-left:
            18px !important;

        padding-right:
            18px !important;
    }


    .page-title {

        font-size:
            31px;
    }

}

</style>
"""
)



# =========================================================
# QUESTION GENERATION LOGS MODAL
# =========================================================

@st.dialog(
    "Question Generation Logs",
    width="large",
)
def question_logs_modal():

    with st.container(key="question_logs_scroll"):

        # =====================================================
        # HEADER
        # =====================================================

        st.html(
        """
        <div class="logs-header">

            <div class="logs-title">
                Question generation logs
            </div>

            <div class="logs-description">
                Search and review previously saved question
                sets. View or edit a record directly below
                the selected row.
            </div>

        </div>
        """
        )


        if st.session_state.get(
            "logs_flash_message"
        ):

            st.success(
                st.session_state[
                    "logs_flash_message"
                ]
            )

            st.session_state.logs_flash_message = None


        # =====================================================
        # LOAD SAVED QUESTION SETS
        # =====================================================

        try:

            question_sets = get_question_sets()

        except Exception as exc:

            print(
                "Question logs error:",
                exc,
            )

            st.error(
                "Unable to load question generation logs. "
                "Please check the backend connection."
            )

            return


        if question_sets is None:
            question_sets = []


        # =====================================================
        # SEARCH
        # =====================================================

        search = st.text_input(
            "Search question sets",
            placeholder="Search by title or topic",
            key="question_logs_search",
        )


        if search:

            search_value = (
                search
                .lower()
                .strip()
            )


            question_sets = [

                item

                for item in question_sets

                if (
                    search_value
                    in str(
                        item.get(
                            "title",
                            "",
                        )
                    ).lower()

                    or

                    search_value
                    in str(
                        item.get(
                            "topic",
                            "",
                        )
                    ).lower()
                )
            ]


        # =====================================================
        # SUMMARY
        # =====================================================

        st.html(
            f"""
            <div class="logs-summary">

                <div class="logs-summary-title">
                    Saved question sets
                </div>

                <div class="logs-summary-text">
                    {len(question_sets)}
                    record(s) match the current search.
                </div>

            </div>
            """
        )


        # =====================================================
        # EMPTY STATE
        # =====================================================

        if not question_sets:

            st.info(
                "No saved question sets are available."
            )

            return


        # =====================================================
        # TABLE HEADER
        # =====================================================

        widths = [
            0.45,
            1.85,
            1.40,
            1.00,
            0.72,
            0.62,
            1.55,
        ]


        header = st.columns(
            widths,
            gap="small",
        )


        headings = [
            "ID",
            "Title",
            "Topic",
            "Difficulty",
            "Questions",
            "Marks",
            "Action",
        ]


        for column, heading in zip(
            header,
            headings,
        ):

            column.html(
                f"""
                <div class="log-table-header">
                    {html.escape(heading)}
                </div>
                """
            )


        # =====================================================
        # TABLE ROWS
        # =====================================================

        for item in question_sets:

            item_id = item.get(
                "id"
            )


            row = st.columns(
                widths,
                gap="small",
            )


            row[0].html(
                f"""
                <div class="log-row-value">
                    {html.escape(str(item_id))}
                </div>
                """
            )


            title = str(
                item.get(
                    "title",
                    "Untitled",
                )
            )


            row[1].html(
                f"""
                <div class="log-row-value">
                    {html.escape(title)}
                </div>
                """
            )


            topic_value = str(
                item.get(
                    "topic",
                    "-",
                )
            )


            row[2].html(
                f"""
                <div class="log-row-value">
                    {html.escape(topic_value)}
                </div>
                """
            )


            difficulty_value = str(
                item.get(
                    "difficulty",
                    "N/A",
                )
            )


            difficulty_class = {

                "Easy":
                    "log-easy",

                "Medium":
                    "log-medium",

                "Hard":
                    "log-hard",

            }.get(
                difficulty_value,
                "log-easy",
            )


            row[3].html(
                f"""
                <span class="log-badge {difficulty_class}">
                    {html.escape(difficulty_value)}
                </span>
                """
            )


            row[4].html(
                f"""
                <div class="log-row-value">
                    {
                        html.escape(
                            str(
                                item.get(
                                    "total_questions",
                                    0,
                                )
                            )
                        )
                    }
                </div>
                """
            )


            displayed_total_marks = (
                st.session_state.get(
                    "question_set_total_overrides",
                    {},
                ).get(
                    item_id,
                    item.get(
                        "total_marks",
                        0,
                    ),
                )
            )


            row[5].html(
                f"""
                <div class="log-row-value">
                    {
                        html.escape(
                            str(
                                displayed_total_marks
                            )
                        )
                    }
                </div>
                """
            )


            with row[6]:

                action_view_col, action_edit_col = st.columns(
                    2,
                    gap="small",
                )


                with action_view_col:

                    if st.button(
                        "View",
                        key=f"log_view_{item_id}",
                        use_container_width=True,
                    ):

                        try:

                            details = get_question_set(
                                item_id
                            )


                            st.session_state.selected_log_set = (
                                normalize_question_set_details(
                                    details,
                                    fallback_id=item_id,
                                )
                            )


                            st.session_state.selected_log_set_id = (
                                item_id
                            )


                            st.session_state.question_log_mode = (
                                "view"
                            )

                            st.session_state.editing_question_id = (
                                None
                            )

                            st.session_state.confirm_delete_log_id = (
                                None
                            )

                        except Exception as exc:

                            print(
                                "Open question set error:",
                                exc,
                            )

                            st.error(
                                "Unable to open this question set."
                            )


                with action_edit_col:

                    if st.button(
                        "Edit",
                        key=f"log_edit_{item_id}",
                        use_container_width=True,
                    ):

                        try:

                            details = get_question_set(
                                item_id
                            )


                            st.session_state.selected_log_set = (
                                normalize_question_set_details(
                                    details,
                                    fallback_id=item_id,
                                )
                            )


                            st.session_state.selected_log_set_id = (
                                item_id
                            )


                            st.session_state.question_log_mode = (
                                "edit"
                            )

                            st.session_state.editing_question_id = (
                                None
                            )

                            st.session_state.confirm_delete_log_id = (
                                None
                            )

                        except Exception as exc:

                            print(
                                "Edit question set error:",
                                exc,
                            )

                            st.error(
                                "Unable to load this question set for editing."
                            )


            st.html(
            """
            <div class="log-divider"></div>
            """
            )



            # =====================================================
            # INLINE VIEW / EDIT EXPANDER FOR THIS ROW
            # =====================================================

            if (
                st.session_state.get(
                    "selected_log_set_id"
                )
                == item_id
            ):

                # =====================================================
                # SELECTED QUESTION SET
                # =====================================================

                selected = st.session_state.get(
                    "selected_log_set"
                )


                if not selected:
                    return


                selected_id = (
                    selected.get(
                        "id"
                    )
                    or
                    selected.get(
                        "set_id"
                    )
                    or
                    st.session_state.get(
                        "selected_log_set_id"
                    )
                )


                selected_title = str(
                    selected.get(
                        "title",
                        "Untitled Question Set",
                    )
                )


                selected_topic = str(
                    selected.get(
                        "topic",
                        "-",
                    )
                )


                selected_difficulty = str(
                    selected.get(
                        "difficulty",
                        "-",
                    )
                )


                selected_questions_for_total = selected.get(
                    "questions",
                    [],
                )


                if selected_questions_for_total:

                    selected_marks = calculate_total_marks_from_questions(
                        selected_questions_for_total
                    )

                else:

                    selected_marks = selected.get(
                        "total_marks",
                        0,
                    )


                panel_mode = st.session_state.get(
                    "question_log_mode",
                    "view",
                )


                panel_action = (
                    "Edit"
                    if panel_mode == "edit"
                    else "View"
                )


                panel_label = (
                    f"{panel_action} — {selected_title}"
                )


                with st.expander(
                    panel_label,
                    expanded=True,
                ):

                    # =====================================================
                    # EDIT MODE
                    # =====================================================

                    if (
                        st.session_state.get(
                            "question_log_mode",
                            "view",
                        )
                        == "edit"
                    ):

                        st.html(
                            f"""
                            <div class="edit-panel-header">

                                <div class="edit-panel-title">
                                    Edit question set
                                </div>

                                <div class="edit-panel-text">
                                    Editing:
                                    <strong>
                                        {html.escape(selected_title)}
                                    </strong>

                                    &nbsp;&nbsp; | &nbsp;&nbsp;

                                    Topic:
                                    {html.escape(selected_topic)}

                                    <br>

                                    Edit the questions below. Save Changes and
                                    Close Edit are available at the end of this section.
                                </div>

                            </div>
                            """
                        )


                        edit_questions = selected.get(
                            "questions",
                            [],
                        )


                        if not edit_questions:

                            st.info(
                                "No questions are available to edit."
                            )

                            return


                        # =================================================
                        # COLLECT ALL QUESTION EDITS
                        # =================================================

                        pending_question_updates = []


                        for index, question in enumerate(
                            edit_questions,
                            start=1,
                        ):

                            question_id = question.get(
                                "id"
                            )


                            question_number = question.get(
                                "question_number",
                                index,
                            )


                            current_text = str(
                                question.get(
                                    "question_text",
                                    question.get(
                                        "question",
                                        "",
                                    ),
                                )
                            )


                            current_type = str(
                                question.get(
                                    "type",
                                    "MCQ",
                                )
                            )


                            current_marks = int(
                                question.get(
                                    "marks",
                                    0,
                                )
                                or 0
                            )


                            current_options = normalize_options(
                                question.get(
                                    "options"
                                )
                            )


                            # Keep the existing correct answer internally.
                            # The field is intentionally not shown on the edit screen.
                            current_answer = (
                                question.get(
                                    "correct_answer"
                                )
                            )


                            with st.container(
                                border=True,
                            ):

                                st.html(
                                    f"""
                                    <div class="edit-question-label">
                                        Question {
                                            html.escape(
                                                str(
                                                    question_number
                                                )
                                            )
                                        }
                                    </div>
                                    """
                                )


                                edited_text = st.text_area(
                                    "Question Text",
                                    value=current_text,
                                    height=120,
                                    key=f"edit_question_text_{question_id}",
                                )


                                type_col, marks_col = st.columns(
                                    2,
                                    gap="large",
                                )


                                question_types = [
                                    "MCQ",
                                    "Short",
                                    "Long",
                                    "Coding",
                                    "Debugging",
                                    "Scenario",
                                    "Case Study",
                                    "Mini Project",
                                    "Capstone",
                                ]


                                if current_type not in question_types:

                                    question_types.insert(
                                        0,
                                        current_type,
                                    )


                                with type_col:

                                    edited_type = st.selectbox(
                                        "Question Type",
                                        question_types,
                                        index=question_types.index(
                                            current_type
                                        ),
                                        key=f"edit_question_type_{question_id}",
                                    )


                                with marks_col:

                                    edited_marks = st.number_input(
                                        "Marks",
                                        min_value=0,
                                        value=current_marks,
                                        step=1,
                                        key=f"edit_question_marks_{question_id}",
                                    )


                                edited_options_text = st.text_area(
                                    "Options",
                                    value="\n".join(
                                        current_options
                                    ),
                                    placeholder=(
                                        "Enter one option per line. "
                                        "Leave empty for non-MCQ questions."
                                    ),
                                    height=110,
                                    key=f"edit_question_options_{question_id}",
                                )


                                edited_options = [
                                    option.strip()

                                    for option
                                    in edited_options_text.splitlines()

                                    if option.strip()
                                ]


                                pending_question_updates.append(
                                    {
                                        "question_id":
                                            question_id,

                                        "question_number":
                                            question_number,

                                        "question_text":
                                            edited_text.strip(),

                                        "payload": {
                                            "question_text":
                                                edited_text.strip(),

                                            "type":
                                                edited_type,

                                            "options":
                                                edited_options,

                                            # Preserve the current value without
                                            # showing an editable Correct Answer field.
                                            "correct_answer":
                                                current_answer,

                                            "marks":
                                                int(
                                                    edited_marks
                                                ),
                                        },
                                    }
                                )


                        # =================================================
                        # EDITED TOTAL MARKS SUMMARY
                        # =================================================

                        edited_total_marks = sum(
                            int(
                                item[
                                    "payload"
                                ][
                                    "marks"
                                ]
                            )

                            for item
                            in pending_question_updates
                        )


                        st.html(
                            f"""
                            <div class="logs-summary">

                                <div class="logs-summary-title">
                                    Updated total marks
                                </div>

                                <div class="logs-summary-text">
                                    The total marks after these edits will be
                                    <strong>{edited_total_marks}</strong>.
                                    This is calculated automatically from all
                                    individual question marks.
                                </div>

                            </div>
                            """
                        )


                        # =================================================
                        # ONE SAVE BUTTON FOR THE COMPLETE QUESTION SET
                        # =================================================

                        if st.button(
                            "Save Changes",
                            key=f"save_question_edit_all_{selected_id}",
                            use_container_width=True,
                        ):

                            # =============================================
                            # VALIDATE ALL QUESTIONS FIRST
                            # =============================================

                            missing_ids = [
                                str(
                                    item[
                                        "question_number"
                                    ]
                                )

                                for item
                                in pending_question_updates

                                if not item[
                                    "question_id"
                                ]
                            ]


                            empty_questions = [
                                str(
                                    item[
                                        "question_number"
                                    ]
                                )

                                for item
                                in pending_question_updates

                                if not item[
                                    "question_text"
                                ]
                            ]


                            if missing_ids:

                                st.error(
                                    "Question ID is missing for question(s): "
                                    + ", ".join(
                                        missing_ids
                                    )
                                    + ". These questions cannot be updated."
                                )


                            elif empty_questions:

                                st.warning(
                                    "Question text cannot be empty for question(s): "
                                    + ", ".join(
                                        empty_questions
                                    )
                                    + "."
                                )


                            else:

                                # =========================================
                                # UPDATE ALL QUESTIONS
                                # =========================================

                                failed_updates = []


                                for item in pending_question_updates:

                                    try:

                                        update_question(
                                            item[
                                                "question_id"
                                            ],
                                            item[
                                                "payload"
                                            ],
                                        )


                                    except Exception as exc:

                                        print(
                                            (
                                                "Update question "
                                                f"{item['question_number']} error:"
                                            ),
                                            exc,
                                        )


                                        failed_updates.append(
                                            {
                                                "question_number":
                                                    item[
                                                        "question_number"
                                                    ],

                                                "error":
                                                    str(
                                                        exc
                                                    ),
                                            }
                                        )


                                # =========================================
                                # REFRESH QUESTION SET AFTER SAVE
                                # =========================================

                                try:

                                    refreshed_data = (
                                        get_question_set(
                                            selected_id
                                        )
                                    )


                                    refreshed_set = (
                                        normalize_question_set_details(
                                            refreshed_data,
                                            fallback_id=selected_id,
                                        )
                                    )


                                    st.session_state.selected_log_set = (
                                        refreshed_set
                                    )


                                    st.session_state.selected_log_set_id = (
                                        selected_id
                                    )


                                    recalculated_total = (
                                        calculate_total_marks_from_questions(
                                            refreshed_set.get(
                                                "questions",
                                                [],
                                            )
                                        )
                                    )


                                    st.session_state[
                                        "question_set_total_overrides"
                                    ][
                                        selected_id
                                    ] = recalculated_total


                                except Exception as refresh_error:

                                    print(
                                        "Refresh question set error:",
                                        refresh_error,
                                    )


                                    # Use the current edited marks in the UI
                                    # even if the refresh request fails.
                                    st.session_state[
                                        "question_set_total_overrides"
                                    ][
                                        selected_id
                                    ] = edited_total_marks


                                # =========================================
                                # FINAL RESULT MESSAGE
                                # =========================================

                                if failed_updates:

                                    failed_numbers = ", ".join(
                                        str(
                                            item[
                                                "question_number"
                                            ]
                                        )

                                        for item
                                        in failed_updates
                                    )


                                    st.error(
                                        "Some questions could not be updated. "
                                        f"Failed question(s): {failed_numbers}."
                                    )


                                else:

                                    st.session_state.logs_flash_message = (
                                        "All question changes were saved successfully. "
                                        f"Total marks: {edited_total_marks}."
                                    )

                                    st.session_state.selected_log_set = None

                                    st.session_state.selected_log_set_id = None

                                    st.session_state.question_log_mode = (
                                        "view"
                                    )

                                    st.session_state.editing_question_id = (
                                        None
                                    )

                                    st.session_state.confirm_delete_log_id = (
                                        None
                                    )

                                    st.rerun(
                                        scope="fragment"
                                    )


                        # =================================================
                        # CLOSE EDIT
                        # Keep this at the bottom for consistency with View.
                        # =================================================

                        if st.button(
                            "Close Edit",
                            key=f"close_log_edit_{selected_id}",
                            use_container_width=False,
                        ):

                            st.session_state.selected_log_set = None

                            st.session_state.selected_log_set_id = None

                            st.session_state.question_log_mode = (
                                "view"
                            )

                            st.session_state.editing_question_id = (
                                None
                            )

                            st.session_state.confirm_delete_log_id = (
                                None
                            )

                            st.rerun(
                                scope="fragment"
                            )


                        return


                    # =====================================================
                    # DETAIL HEADER
                    # =====================================================

                    st.html(
                        f"""
                        <div class="log-detail">

                            <div class="log-detail-title">
                                {html.escape(selected_title)}
                            </div>

                            <div class="log-detail-meta">

                                Topic:
                                {html.escape(selected_topic)}

                                &nbsp;&nbsp; | &nbsp;&nbsp;

                                Difficulty:
                                {html.escape(selected_difficulty)}

                                &nbsp;&nbsp; | &nbsp;&nbsp;

                                Total marks:
                                {html.escape(str(selected_marks))}

                            </div>

                        </div>
                        """
                    )


                    # =====================================================
                    # QUESTIONS
                    # =====================================================

                    questions = selected.get(
                        "questions",
                        [],
                    )


                    if not questions:

                        st.info(
                            "No questions are available for this question set."
                        )


                    else:

                        for index, question in enumerate(
                            questions,
                            start=1,
                        ):

                            question_number = question.get(
                                "question_number",
                                index,
                            )


                            question_text = question.get(
                                "question_text",
                                question.get(
                                    "question",
                                    "",
                                ),
                            )


                            question_type_value = question.get(
                                "type",
                                "-",
                            )


                            question_marks = question.get(
                                "marks",
                                0,
                            )


                            st.html(
                                f"""
                                <div class="log-question">

                                    <div class="log-question-number">
                                        Question {
                                            html.escape(
                                                str(
                                                    question_number
                                                )
                                            )
                                        }
                                    </div>

                                    <div class="log-question-text">
                                        {
                                            html.escape(
                                                str(
                                                    question_text
                                                )
                                            )
                                        }
                                    </div>

                                    <div class="log-question-meta">

                                        Type:
                                        {
                                            html.escape(
                                                str(
                                                    question_type_value
                                                )
                                            )
                                        }

                                        &nbsp;&nbsp; | &nbsp;&nbsp;

                                        Marks:
                                        {
                                            html.escape(
                                                str(
                                                    question_marks
                                                )
                                            )
                                        }

                                    </div>

                                </div>
                                """
                            )


                            options = question.get(
                                "options"
                            )


                            if options:

                                with st.expander(
                                    f"Options for Question {question_number}"
                                ):

                                    for option in options:
                                        st.write(option)


                    # =====================================================
                    # PREPARE RELIABLE DOWNLOAD FILES
                    # Backend export is preferred.
                    # Local generation is used automatically as fallback.
                    # =====================================================

                    try:

                        pdf_data = get_reliable_pdf_export(
                            selected_id,
                            selected,
                        )

                    except Exception as exc:

                        print(
                            "Question log PDF generation error:",
                            exc,
                        )

                        pdf_data = None


                    try:

                        docx_data = get_reliable_docx_export(
                            selected_id,
                            selected,
                        )

                    except Exception as exc:

                        print(
                            "Question log DOCX generation error:",
                            exc,
                        )

                        docx_data = None


                    # =====================================================
                    # ACTIONS
                    # =====================================================

                    action_col1, action_col2, action_col3 = st.columns(
                        3,
                        gap="large",
                    )


                    with action_col1:

                        if pdf_data is not None:

                            st.download_button(
                                label="Download PDF",
                                data=pdf_data,
                                file_name=f"EduAI_Question_Set_{selected_id}.pdf",
                                mime="application/pdf",
                                key=f"log_pdf_{selected_id}",
                                use_container_width=True,
                                on_click="ignore",
                            )

                        else:

                            st.button(
                                "PDF generation failed",
                                disabled=True,
                                key=f"log_pdf_disabled_{selected_id}",
                                use_container_width=True,
                            )


                    with action_col2:

                        if docx_data is not None:

                            st.download_button(
                                label="Download DOCX",
                                data=docx_data,
                                file_name=f"EduAI_Question_Set_{selected_id}.docx",
                                mime=(
                                    "application/vnd.openxmlformats-officedocument."
                                    "wordprocessingml.document"
                                ),
                                key=f"log_docx_{selected_id}",
                                use_container_width=True,
                                on_click="ignore",
                            )

                        else:

                            st.button(
                                "DOCX generation failed",
                                disabled=True,
                                key=f"log_docx_disabled_{selected_id}",
                                use_container_width=True,
                            )


                    with action_col3:

                        if st.button(
                            "Delete Question Set",
                            key=f"log_delete_{selected_id}",
                            use_container_width=True,
                        ):

                            st.session_state.confirm_delete_log_id = (
                                selected_id
                            )


                    # =====================================================
                    # DELETE CONFIRMATION
                    # =====================================================

                    if (
                        st.session_state.get(
                            "confirm_delete_log_id"
                        )
                        == selected_id
                    ):

                        st.warning(
                            "This question set will be permanently deleted."
                        )


                        confirm_col1, confirm_col2 = st.columns(
                            2,
                            gap="large",
                        )


                        with confirm_col1:

                            if st.button(
                                "Confirm Delete",
                                key=f"log_confirm_delete_{selected_id}",
                                use_container_width=True,
                            ):

                                try:

                                    delete_question_set(
                                        selected_id
                                    )


                                    st.session_state.selected_log_set = None

                                    st.session_state.selected_log_set_id = None

                                    st.session_state.confirm_delete_log_id = None


                                    st.session_state.get(
                                        "question_set_total_overrides",
                                        {},
                                    ).pop(
                                        selected_id,
                                        None,
                                    )


                                    st.success(
                                        "Question set deleted successfully."
                                    )


                                except Exception as exc:

                                    print(
                                        "Delete question set error:",
                                        exc,
                                    )

                                    st.error(
                                        "Unable to delete the question set."
                                    )


                        with confirm_col2:

                            if st.button(
                                "Cancel",
                                key=f"log_cancel_delete_{selected_id}",
                                use_container_width=True,
                            ):

                                st.session_state.confirm_delete_log_id = None

                    # =====================================================
                    # CLOSE VIEW
                    # =====================================================

                    if st.button(
                        "Close View",
                        key=f"close_log_view_{selected_id}",
                        use_container_width=False,
                    ):

                        st.session_state.selected_log_set = None

                        st.session_state.selected_log_set_id = None

                        st.session_state.confirm_delete_log_id = None

                        st.session_state.question_log_mode = (
                            "view"
                        )

                        st.rerun(
                            scope="fragment"
                        )




# =========================================================
# GENERATION MODAL
# =========================================================

@st.dialog(
    "Generate Assessment",
    width="large",
)
def generation_modal():

    with st.container(key="generation_modal_scroll"):

        snapshot = st.session_state.get(
            "generation_snapshot"
        )


        # =====================================================
        # CONFIG MISSING
        # =====================================================

        if not snapshot:

            st.html(
            """
            <div class="error-card">

                <div class="error-title">
                    Configuration unavailable
                </div>

                <div class="error-text">
                    Assessment configuration could not be loaded.
                    Close this window and try again.
                </div>

            </div>
            """
            )

            return


        # =====================================================
        # HEADER
        # =====================================================

        st.html(
        """
        <div class="modal-header">

            <div class="modal-title">
                Assessment generation
            </div>

            <div class="modal-description">
                Question generation, formatted preview,
                saving and downloading are handled inside
                this window.
            </div>

        </div>
        """
        )


        # =====================================================
        # CONFIGURATION SUMMARY
        # =====================================================

        st.html(
            f"""
            <div class="modal-config">

                <strong>Page format:</strong>
                {html.escape(snapshot["letterhead_option"])}

                &nbsp;&nbsp; | &nbsp;&nbsp;

                <strong>Source:</strong>
                {html.escape(snapshot["mode"])}

                &nbsp;&nbsp; | &nbsp;&nbsp;

                <strong>Question type:</strong>
                {html.escape(snapshot["question_type"])}

                &nbsp;&nbsp; | &nbsp;&nbsp;

                <strong>Difficulty:</strong>
                {html.escape(snapshot["difficulty"])}

                &nbsp;&nbsp; | &nbsp;&nbsp;

                <strong>Questions:</strong>
                {snapshot["number_of_questions"]}

                &nbsp;&nbsp; | &nbsp;&nbsp;

                <strong>Total marks:</strong>
                {snapshot["total_marks"]}

            </div>
            """
        )


        phase = st.session_state.get(
            "generation_modal_phase",
            "generate",
        )


        # =====================================================
        # GENERATE
        # =====================================================

        if phase == "generate":

            loading_area = st.empty()


            loading_area.html(
            """
            <div class="loading-card">

                <div class="loading-title">
                    Generating your assessment
                </div>

                <div class="loading-text">
                    The selected source is being analyzed and
                    your question type, difficulty level, marks
                    and assessment requirements are being applied.
                </div>

                <div class="loading-track">
                    <div class="loading-bar"></div>
                </div>

                <div class="loading-note">
                    Generation is in progress.
                    Please keep this window open.
                </div>

            </div>
            """
            )


            # =================================================
            # VALIDATION - LETTERHEAD
            # =================================================

            if (
                snapshot["letterhead_option"]
                == "Upload Letterhead"
                and not snapshot.get(
                    "letterhead_file"
                )
            ):

                loading_area.empty()

                st.session_state.generation_modal_phase = (
                    "error"
                )

                st.session_state.generation_error = (
                    "Please upload a letterhead file "
                    "before generating the assessment."
                )

                phase = "error"


            # =================================================
            # VALIDATION - CUSTOM FORMAT
            # =================================================

            elif (
                snapshot["letterhead_option"]
                == "Custom Instructions"
                and not snapshot.get(
                    "letterhead_instruction",
                    "",
                ).strip()
            ):

                loading_area.empty()

                st.session_state.generation_modal_phase = (
                    "error"
                )

                st.session_state.generation_error = (
                    "Please provide page format instructions "
                    "before generating the assessment."
                )

                phase = "error"


            # =================================================
            # VALIDATION - TOPIC
            # =================================================

            elif (
                snapshot["mode"]
                == "Topic Based"
                and not snapshot.get(
                    "topic",
                    "",
                ).strip()
            ):

                loading_area.empty()

                st.session_state.generation_modal_phase = (
                    "error"
                )

                st.session_state.generation_error = (
                    "Please enter a topic before "
                    "generating the assessment."
                )

                phase = "error"


            # =================================================
            # VALIDATION - MATERIAL
            # =================================================

            elif (
                snapshot["mode"]
                == "Upload Material"
                and not snapshot.get(
                    "learning_file"
                )
            ):

                loading_area.empty()

                st.session_state.generation_modal_phase = (
                    "error"
                )

                st.session_state.generation_error = (
                    "Please upload learning material before "
                    "generating the assessment."
                )

                phase = "error"


            # =================================================
            # RUN GENERATION
            # =================================================

            else:

                try:

                    # =========================================
                    # TOPIC BASED
                    # =========================================

                    if (
                        snapshot["mode"]
                        == "Topic Based"
                    ):

                        payload = {

                            "source_type":
                                "topic",

                            "topic":
                                snapshot[
                                    "topic"
                                ].strip(),

                            "material_content":
                                None,

                            "question_type":
                                snapshot[
                                    "question_type"
                                ],

                            "difficulty":
                                snapshot[
                                    "difficulty"
                                ],

                            "number_of_questions":
                                snapshot[
                                    "number_of_questions"
                                ],

                            "total_marks":
                                snapshot[
                                    "total_marks"
                                ],
                        }


                        result = generate_questions(
                            payload
                        )


                    # =========================================
                    # MATERIAL BASED
                    # =========================================

                    else:

                        learning_file = rebuild_file(
                            snapshot[
                                "learning_file"
                            ]
                        )


                        result = (
                            generate_questions_from_file(
                                learning_file,
                                snapshot[
                                    "question_type"
                                ],
                                snapshot[
                                    "difficulty"
                                ],
                                snapshot[
                                    "number_of_questions"
                                ],
                                snapshot[
                                    "total_marks"
                                ],
                            )
                        )


                    # =========================================
                    # RESULT METADATA
                    # =========================================

                    result["difficulty"] = (
                        snapshot[
                            "difficulty"
                        ]
                    )


                    result["total_questions"] = (
                        snapshot[
                            "number_of_questions"
                        ]
                    )


                    result["total_marks"] = (
                        snapshot[
                            "total_marks"
                        ]
                    )


                    # =========================================
                    # STORE RESULT
                    # =========================================

                    st.session_state.generated_questions = (
                        result
                    )


                    st.session_state.preview_pdf = None


                    st.session_state.generation_error = None


                    st.session_state.save_success_message = None


                    st.session_state.generation_modal_phase = (
                        "generated"
                    )


                    phase = "generated"


                    loading_area.empty()


                except Exception as exc:

                    print(
                        "Question generation error:",
                        exc,
                    )


                    loading_area.empty()


                    st.session_state.generation_modal_phase = (
                        "error"
                    )


                    st.session_state.generation_error = (
                        "Unable to generate questions right now. "
                        "Please verify the Question Generation "
                        "service and try again."
                    )


                    phase = "error"


        # =====================================================
        # ERROR
        # =====================================================

        if phase == "error":

            error_message = (
                st.session_state.get(
                    "generation_error"
                )
                or
                "Unable to generate the assessment."
            )


            st.html(
                f"""
                <div class="error-card">

                    <div class="error-title">
                        Assessment could not be generated
                    </div>

                    <div class="error-text">
                        {html.escape(error_message)}
                    </div>

                </div>
                """
            )


            st.caption(
                "Close this window, update the required "
                "configuration and try again."
            )


            return


        # =====================================================
        # GENERATED
        # =====================================================

        if phase == "generated":

            result = st.session_state.get(
                "generated_questions"
            )


            if result:

                question_count = len(
                    result.get(
                        "questions",
                        [],
                    )
                )

            else:

                question_count = 0


            st.html(
                f"""
                <div class="success-card">

                    <div class="success-title">
                        Assessment generated successfully
                    </div>

                    <div class="success-text">

                        {question_count} question(s)
                        have been generated successfully.

                        Review the formatted question paper
                        before saving or downloading it.

                    </div>

                </div>
                """
            )


            if st.button(
                "Preview Question Paper",
                key="modal_preview_question_paper",
                use_container_width=True,
            ):

                st.session_state.generation_modal_phase = (
                    "preview"
                )

                phase = "preview"


        # =====================================================
        # PREVIEW
        # =====================================================

        if phase == "preview":

            data = st.session_state.get(
                "generated_questions"
            )


            if not data:

                st.html(
                """
                <div class="error-card">

                    <div class="error-title">
                        Preview unavailable
                    </div>

                    <div class="error-text">
                        Generated questions could not be found.
                    </div>

                </div>
                """
                )

                return


            # =================================================
            # PREVIEW HEADING
            # =================================================

            st.html(
            """
            <div class="modal-header">

                <div class="modal-title">
                    Question paper preview
                </div>

                <div class="modal-description">
                    Review the final paper exactly in the
                    selected page format before saving
                    or downloading.
                </div>

            </div>
            """
            )


            # =================================================
            # CREATE PDF
            # =================================================

            if st.session_state.preview_pdf is None:

                pdf_loading = st.empty()


                pdf_loading.html(
                """
                <div class="loading-card">

                    <div class="loading-title">
                        Preparing question paper preview
                    </div>

                    <div class="loading-text">
                        The selected page format is being
                        applied to the generated assessment.
                    </div>

                    <div class="loading-track">
                        <div class="loading-bar"></div>
                    </div>

                    <div class="loading-note">
                        Please wait while the formatted paper
                        is prepared.
                    </div>

                </div>
                """
                )


                try:

                    letterhead_file = rebuild_file(
                        snapshot.get(
                            "letterhead_file"
                        )
                    )


                    pdf_file = (
                        generate_question_paper_pdf(
                            data=data,

                            letterhead_option=
                                snapshot[
                                    "letterhead_option"
                                ],

                            letterhead_file=
                                letterhead_file,

                            letterhead_instruction=
                                snapshot.get(
                                    "letterhead_instruction"
                                ),
                        )
                    )


                    pdf_bytes = get_pdf_bytes(
                        pdf_file
                    )


                    st.session_state.preview_pdf = (
                        pdf_bytes
                    )


                    pdf_loading.empty()


                except Exception as exc:

                    print(
                        "PDF preview error:",
                        exc,
                    )


                    pdf_loading.empty()


                    st.html(
                    """
                    <div class="error-card">

                        <div class="error-title">
                            Preview could not be prepared
                        </div>

                        <div class="error-text">
                            The formatted question paper
                            could not be created.
                            Please check the selected page
                            format and try again.
                        </div>

                    </div>
                    """
                    )


                    return


            # =================================================
            # PDF DATA
            # =================================================

            pdf_bytes = (
                st.session_state.preview_pdf
            )


            # =================================================
            # PDF LABEL
            # =================================================

            st.html(
                f"""
                <div class="pdf-preview-note">

                    Page format:
                    <strong>
                        {
                            html.escape(
                                snapshot[
                                    "letterhead_option"
                                ]
                            )
                        }
                    </strong>

                </div>
                """
            )


            st.html(
            """
            <div class="pdf-preview-label">
                Formatted question paper
            </div>
            """
            )


            # =================================================
            # NATIVE STREAMLIT PDF VIEWER
            # =================================================

            try:

                st.pdf(
                    pdf_bytes,
                    height=620,
                    key="question_paper_pdf_preview",
                )


            except Exception as exc:

                print(
                    "Native PDF viewer error:",
                    exc,
                )


                st.error(
                    "The PDF was created successfully, but "
                    "the PDF preview component is not available. "
                    'Install it with: pip install "streamlit[pdf]"'
                )


            # =================================================
            # FINAL ACTIONS
            # =================================================

            st.html(
            """
            <div class="modal-actions-label">
                Final actions
            </div>
            """
            )


            action_col1, action_col2 = st.columns(
                2,
                gap="large",
            )


            # =================================================
            # SAVE QUESTION SET
            # =================================================

            with action_col1:

                if st.button(
                    "Save Question Set",
                    key="modal_save_question_set",
                    use_container_width=True,
                ):

                    try:

                        response = save_question_set(
                            data
                        )


                        if isinstance(
                            response,
                            dict,
                        ):

                            message = response.get(
                                "message",
                                "Question set saved successfully.",
                            )

                        else:

                            message = (
                                "Question set saved successfully."
                            )


                        st.session_state.save_success_message = (
                            message
                        )


                    except Exception as exc:

                        print(
                            "Save question set error:",
                            exc,
                        )


                        st.error(
                            "Unable to save the question set. "
                            "Please try again."
                        )


            # =================================================
            # DOWNLOAD
            # =================================================

            with action_col2:

                st.download_button(
                    label="Download Question Paper",

                    data=pdf_bytes,

                    file_name=
                        "EduAI_Question_Paper.pdf",

                    mime=
                        "application/pdf",

                    key=
                        "modal_download_question_paper",

                    use_container_width=True,

                    on_click="ignore",
                )


            # =================================================
            # SAVE SUCCESS
            # =================================================

            if st.session_state.get(
                "save_success_message"
            ):

                st.success(
                    st.session_state[
                        "save_success_message"
                    ]
                )



# =========================================================
# MAIN PAGE HEADER
# =========================================================

back_col, header_col = st.columns(
    [
        0.55,
        9.45,
    ],
    gap="small",
    vertical_alignment="top",
)


with back_col:

    if st.button(
        "‹",
        key="back_to_dashboard",
        help="Back to dashboard",
    ):
        st.switch_page(
            "pages/instructor_dashboard.py"
        )


with header_col:

    st.html(
    """
    <div class="page-header">

        <div class="page-label">
            Question Generation
        </div>

        <div class="page-title">
            Create your assessment
        </div>

        <div class="page-description">
            Select the page format, configure the assessment
            requirements and provide the source content.
            Once generation starts, generation, preview,
            saving and downloading are handled inside
            one focused window.
        </div>

    </div>
    """
    )


# =========================================================
# STEP 01
# PAGE FORMAT
# =========================================================

with st.container(
    border=True,
):

    st.html(
    """
    <div class="section-header">

        <div class="section-number">
            Step 01
        </div>

        <div class="section-title">
            Page format
        </div>

        <div class="section-description">
            Choose how the final question paper should
            be formatted before configuring the assessment.
        </div>

    </div>
    """
    )


    letterhead_option = st.radio(
        "Page Format",

        [
            "Default EduAI Format",
            "Upload Letterhead",
            "Custom Instructions",
        ],

        horizontal=True,

        key="page_format",

        label_visibility="collapsed",
    )


    letterhead_file = None

    letterhead_instruction = None


    # =====================================================
    # UPLOAD LETTERHEAD
    # =====================================================

    if letterhead_option == "Upload Letterhead":

        st.write("")


        letterhead_file = st.file_uploader(
            "Upload Letterhead File",

            type=[
                "pdf",
                "png",
                "jpg",
                "jpeg",
            ],

            key="letterhead_file",
        )


    # =====================================================
    # CUSTOM FORMAT
    # =====================================================

    elif (
        letterhead_option
        == "Custom Instructions"
    ):

        st.write("")


        letterhead_instruction = st.text_area(
            "Page Format Instructions",

            placeholder=(
                "Example:\n\n"
                "University name centered at the top.\n"
                "Department name below the university name.\n"
                "Course information below the header.\n"
                "Use a clean horizontal separator.\n"
                "Add page numbers in the footer."
            ),

            height=160,

            key="letterhead_instruction",
        )


# =========================================================
# STEP 02
# QUESTION SETTINGS
# =========================================================

with st.container(
    border=True,
):

    # =====================================================
    # TESTED HEADER LAYOUT:
    # LEFT TITLE + RIGHT VIEW LOGS BUTTON
    # =====================================================

    header_left, header_space, header_right = st.columns(
        [6.5, 2.0, 1.5],
        gap="small",
    )


    with header_left:

        st.html(
        """
        <div class="section-header">

            <div class="section-number">
                Step 02
            </div>

            <div class="section-title">
                Question generation
            </div>

            <div class="section-description">
                Configure the source, question type,
                difficulty level, quantity and total marks.
            </div>

        </div>
        """
        )


    with header_right:

        st.write("")


        if st.button(
            "Generated Records",
            key="view_logs_button",
            use_container_width=True,
        ):

            question_logs_modal()


    # =====================================================
    # GENERATION SOURCE
    # =====================================================

    mode = st.radio(
        "Generation Source",

        [
            "Topic Based",
            "Upload Material",
        ],

        horizontal=True,

        key="generation_source",

        label_visibility="collapsed",
    )


    st.write("")


    # =====================================================
    # TYPE + DIFFICULTY
    # =====================================================

    config_col1, config_col2 = st.columns(
        2,
        gap="large",
    )


    with config_col1:

        question_type = st.selectbox(
            "Question Type",

            [
                "MCQ",
                "Short",
                "Long",
                "Coding",
                "Debugging",
                "Scenario",
                "Case Study",
                "Mini Project",
                "Capstone",
            ],

            key="question_type",
        )


    with config_col2:

        difficulty = st.selectbox(
            "Difficulty",

            [
                "Easy",
                "Medium",
                "Hard",
            ],

            key="difficulty",
        )


    # =====================================================
    # NUMBER + MARKS
    # =====================================================

    config_col3, config_col4 = st.columns(
        2,
        gap="large",
    )


    with config_col3:

        number_of_questions = st.number_input(
            "Number of Questions",

            min_value=1,

            value=5,

            step=1,

            key="number_of_questions",
        )


    with config_col4:

        total_marks = st.number_input(
            "Total Marks",

            min_value=1,

            value=10,

            step=1,

            key="total_marks",
        )


# =========================================================
# STEP 03
# SOURCE CONTENT
# =========================================================

topic = None

uploaded_file = None


with st.container(
    border=True,
):

    # =====================================================
    # TOPIC
    # =====================================================

    if mode == "Topic Based":

        st.html(
        """
        <div class="section-header">

            <div class="section-number">
                Step 03
            </div>

            <div class="section-title">
                Topic information
            </div>

            <div class="section-description">
                Enter the topic that should be used as
                the basis for question generation.
            </div>

        </div>
        """
        )


        topic = st.text_input(
            "Topic",

            placeholder=
                "Example: Artificial Intelligence Agents",

            key=
                "topic_input",
        )


    # =====================================================
    # MATERIAL
    # =====================================================

    else:

        st.html(
        """
        <div class="section-header">

            <div class="section-number">
                Step 03
            </div>

            <div class="section-title">
                Learning material
            </div>

            <div class="section-description">
                Upload the source material that should
                be analyzed for question generation.
            </div>

        </div>
        """
        )


        uploaded_file = st.file_uploader(
            "Upload PDF, DOCX or PPTX",

            type=[
                "pdf",
                "docx",
                "pptx",
            ],

            key=
                "learning_material",
        )


# =========================================================
# STEP 04
# GENERATE
# =========================================================

with st.container(
    border=True,
):

    st.html(
    """
    <div class="section-header">

        <div class="section-number">
            Step 04
        </div>

        <div class="section-title">
            Generate assessment
        </div>

        <div class="section-description">
            Review the current configuration and start
            generation when everything is ready.
        </div>

    </div>
    """
    )


    source_summary = (
        "Topic Based"
        if mode == "Topic Based"
        else "Uploaded Material"
    )


    # =====================================================
    # SUMMARY
    # =====================================================

    st.html(
        f"""
        <div class="generation-summary">

            <div class="summary-title">
                Current configuration
            </div>

            <div class="summary-text">

                Page format:
                {html.escape(letterhead_option)}

                &nbsp;&nbsp; | &nbsp;&nbsp;

                Source:
                {html.escape(source_summary)}

                &nbsp;&nbsp; | &nbsp;&nbsp;

                Question type:
                {html.escape(question_type)}

                &nbsp;&nbsp; | &nbsp;&nbsp;

                Difficulty:
                {html.escape(difficulty)}

                &nbsp;&nbsp; | &nbsp;&nbsp;

                Questions:
                {int(number_of_questions)}

                &nbsp;&nbsp; | &nbsp;&nbsp;

                Total marks:
                {int(total_marks)}

            </div>

        </div>
        """
    )


    # =====================================================
    # GENERATE BUTTON
    # =====================================================

    if st.button(
        "Generate Questions",
        key="generate_questions_button",
        use_container_width=True,
    ):

        # =================================================
        # SNAPSHOT SETTINGS
        # =================================================

        st.session_state.generation_snapshot = {

            "letterhead_option":
                letterhead_option,

            "letterhead_instruction":
                letterhead_instruction or "",

            "letterhead_file":
                create_file_snapshot(
                    letterhead_file
                ),

            "mode":
                mode,

            "topic":
                topic or "",

            "learning_file":
                create_file_snapshot(
                    uploaded_file
                ),

            "question_type":
                question_type,

            "difficulty":
                difficulty,

            "number_of_questions":
                int(
                    number_of_questions
                ),

            "total_marks":
                int(
                    total_marks
                ),
        }


        # =================================================
        # RESET OLD SESSION RESULTS
        # =================================================

        st.session_state.generated_questions = None

        st.session_state.preview_pdf = None

        st.session_state.generation_error = None

        st.session_state.save_success_message = None

        st.session_state.generation_modal_phase = (
            "generate"
        )


        # =================================================
        # OPEN MODAL
        # =================================================

        generation_modal()